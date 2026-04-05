"""
Run this script on your local machine.
It will:
1. Download all images from your blog
2. Generate the full Word document with images embedded
3. Save it as Chapters_3_4_Final.docx

Requirements:
    pip install python-docx requests Pillow

Run with:
    python build_report.py
"""

import io
import os
import requests
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── IMAGE URLS ──────────────────────────────────────────────────────────────
IMAGES = {
    "fig3_1_rccar":           "https://static.wixstatic.com/media/7531f0_6f8e098556aa4ac3bf74e19aef7ce830~mv2.jpeg",
    "fig3_2_jetson_lidar":    "https://static.wixstatic.com/media/7531f0_43248db75940474989bc1d999d1f90be~mv2.png",
    "fig3_3_pi_radar":        "https://static.wixstatic.com/media/7531f0_345c827e9a8043f98d61abd0b5a1d000~mv2.png",
    "fig3_4_fusion_gui":      "https://static.wixstatic.com/media/7531f0_9c81e24e23a64eee8fce8566ac65e34f~mv2.jpg",
    "fig3_5_lidar_scatter":   "https://static.wixstatic.com/media/7531f0_4f381abe0df4405fa234c331864cfbad~mv2.png",
    "fig3_6_radar_gui":       "https://static.wixstatic.com/media/7531f0_87e1557f52bd4eb29daee4459c72b0ac~mv2.jpg",
    "fig3_7_camera":          "https://static.wixstatic.com/media/7531f0_3621c22f91f3431a91f391852ba8286d~mv2.png",
    "fig3_8_statemachine":    "https://static.wixstatic.com/media/7531f0_b47408112a714655aa4b23d2de4e9015~mv2.png",
    "fig3_9_esc_arming":      "https://static.wixstatic.com/media/7531f0_e71812ae4fd14219a9f82a3d458d7e4e~mv2.png",
    "fig4_1_town04":          "https://static.wixstatic.com/media/7531f0_3b45265d65f7415888d372cd98516336~mv2.png",
    "fig4_2_brake_clear":     "https://static.wixstatic.com/media/7531f0_a7d3d8eef9b14532b75c1ce24caa23c5~mv2.png",
    "fig4_3_lane_change":     "https://static.wixstatic.com/media/7531f0_a7d1dabe23d04a5bb8521790d758b10e~mv2.png",
    "fig4_4_straightening":   "https://static.wixstatic.com/media/7531f0_efc63a95d590466f938bc562166ec3bf~mv2.png",
    "fig4_5_resume":          "https://static.wixstatic.com/media/7531f0_61dbb45a208c492ca6d515f6aedf5db9~mv2.png",
    "fig4_6_dashboard":       "https://static.wixstatic.com/media/7531f0_66dcfbcd2bb94822bf774862e86a6e46~mv2.png",
    "fig4_7_rain_setup":      "https://static.wixstatic.com/media/7531f0_42fd9844f46647ed92c6a2bc7debf693~mv2.png",
    "fig4_8_brake_rain":      "https://static.wixstatic.com/media/7531f0_aeb734b3a5f04be8860281ed049608ee~mv2.png",
    "fig4_9_lanechange_rain": "https://static.wixstatic.com/media/7531f0_1ea52bd54e1e4c80af6e695a6603023c~mv2.png",
    "yolo_v8_baseline":        "https://static.wixstatic.com/media/7531f0_1115d4fffe8f44ac81cdbd466c73f19e~mv2.png",
    "yolo_v5_clear_perf":      "https://static.wixstatic.com/media/7531f0_9ad96f478dcb4f3296474d586cbbc896~mv2.png",
    "yolo_v5_clear_stats":     "https://static.wixstatic.com/media/7531f0_eef513b1e5134cb6bbbde0dd7f550012~mv2.png",
    "yolo_v8_clear_perf":      "https://static.wixstatic.com/media/7531f0_197d02baaeab4fb0b906b62c1307ef5c~mv2.png",
    "yolo_v8_clear_stats":     "https://static.wixstatic.com/media/7531f0_bdbee2fcc5214009a382486b4ec493ba~mv2.png",
    "yolo_v5_rain_perf":       "https://static.wixstatic.com/media/7531f0_97efe3dfa89144a19b5d3cab03e3850b~mv2.png",
    "yolo_v5_rain_stats":      "https://static.wixstatic.com/media/7531f0_48609463559843c1bdb370d088d13976~mv2.png",
    "yolo_v8_rain_perf":       "https://static.wixstatic.com/media/7531f0_15af0a600db84b788fb8ac751a68f65c~mv2.png",
    "yolo_v8_rain_stats":      "https://static.wixstatic.com/media/7531f0_ab5f288419cf4f73a765060a586d6260~mv2.png",
    "fig4_11_lidar_time":     "https://static.wixstatic.com/media/7531f0_b5eee49c9c8f4686818d418a26de1c72~mv2.png",
    "fig4_12_motor_states":   "https://static.wixstatic.com/media/7531f0_17dc20ba7995479a9cc3043e54d38f15~mv2.png",
    "fig4_13_fusion_levels":  "https://static.wixstatic.com/media/7531f0_92759eeee5a84ba785ababfacb15cd1e~mv2.png",
    "fig4_14_yolo_det":       "https://static.wixstatic.com/media/7531f0_c7925ab2982745a28e4bcbe936585433~mv2.png",
    "fig4_15_lidar_hist":     "https://static.wixstatic.com/media/7531f0_5b9bdc6b535042f1a4a47221af691db5~mv2.png",
    "fig4_16_avoidance":      "https://static.wixstatic.com/media/7531f0_e7e083fa94474acea91f6a9f196b9a02~mv2.png",
    "fig4_17_perf":           "https://static.wixstatic.com/media/7531f0_8999efd048b442d2bc0c1b611fde62e7~mv2.png",

}

# ── DOWNLOAD IMAGES ──────────────────────────────────────────────────────────
print("Downloading images from blog...")
img_cache = {}
headers = {"User-Agent": "Mozilla/5.0"}
for key, url in IMAGES.items():
    try:
        r = requests.get(url, headers=headers, timeout=15)
        if r.status_code == 200:
            img_cache[key] = io.BytesIO(r.content)
            print(f"  OK  {key}")
        else:
            print(f"  FAIL {key} ({r.status_code})")
    except Exception as e:
        print(f"  ERROR {key}: {e}")

# ── HELPERS ──────────────────────────────────────────────────────────────────
def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_image(doc, key, caption_text, width=5.5):
    if key in img_cache:
        img_cache[key].seek(0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run()
        run.add_picture(img_cache[key], width=Inches(width))
        add_caption(doc, caption_text)
    else:
        p = doc.add_paragraph(f"[IMAGE NOT DOWNLOADED: {caption_text}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_caption(doc, caption_text)

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)

def add_github(doc, path):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"GitHub: {path}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.italic = True

def add_table(doc, headers, rows, col_widths_inches):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        run = hdr[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        tr = table.add_row().cells
        fill = "EEF3FB" if ri % 2 == 0 else "FFFFFF"
        for i, cell in enumerate(row):
            tr[i].text = cell
            tr[i].paragraphs[0].runs[0].font.size = Pt(9)
            tc = tr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
    # set column widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_inches):
                cell.width = Inches(col_widths_inches[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def section_banner(doc, text, color_hex="2E75B6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16)
    )
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F0F0F0")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)

# ── BUILD DOCUMENT ────────────────────────────────────────────────────────────
doc = Document()

# Set default font and margins
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "Chapter 3: Proposed System Design and Implementation", 1)

add_heading(doc, "3.1 Overview of the Two-Phase Development Approach", 2)
add_para(doc, "The project is structured across two phases with fundamentally different goals. Phase 1 uses the Carla simulator to develop and validate the perception algorithms, sensor fusion logic and TTC computation before any physical hardware is involved. The ego vehicle follows a predefined lane, detects a stationary obstacle through fused sensor data, executes an emergency brake and performs a structured lane change before resuming normal driving. Navigation in Phase 1 is waypoint-guided and planned.")
add_para(doc, "Phase 2 is a physical deployment on a Maverick RC car with no waypoints and no lane maps. The vehicle drives forward by default and reacts to whatever the sensors detect: when an obstacle blocks the path it stops, reverses, turns and escapes on a new heading, and when the path is clear it continues forward. This is a purely reactive architecture appropriate for an embedded platform without localisation or prior environmental knowledge [18].")
add_para(doc, "The Digital Twin layer is present in both phases as the internal software model continuously updated with obstacle positions, distances, velocities and TTC values from live sensor data. In Phase 1 it is populated from Carla simulated sensors and in Phase 2 from real hardware. The architecture connecting sensors to fusion to actuation is the same across both phases, making the transition a change in data source rather than a change in system design [16], [17]. Table 3.1 summarises the key structural differences between the two phases.")

add_table(doc,
    ["Aspect", "Phase 1 Carla Simulation", "Phase 2 Hardware Deployment"],
    [
        ["Platform", "Carla simulator, Town04 map", "Maverick RC car, real environment"],
        ["Navigation", "Waypoint-guided lane following", "Purely reactive, no map"],
        ["Sensors", "Simulated lidar, radar, camera", "LD06, BGT60TR13C, Arducam IMX477"],
        ["Compute", "Laptop running Carla client", "Jetson Orin Nano and Raspberry Pi 5"],
        ["Motor control", "Carla VehicleControl API", "Arduino UDP to PWM, physical ESC and servo"],
        ["Post-obstacle behaviour", "Planned lane change and resume", "Reverse, turn and escape on new heading"],
        ["Obstacle type", "Stationary vehicle at 80 m", "Physical everyday objects, indoor and outdoor"],
    ],
    [2.0, 2.5, 2.5]
)
add_caption(doc, "Table 3.1: Comparison of Phase 1 and Phase 2 system characteristics.")

section_banner(doc, "Phase 1: Simulation Baseline (Carla)", "1A6E3C")

add_heading(doc, "3.2 Phase 1 Simulation Design and Implementation", 2)
add_heading(doc, "3.2.1 Simulation Environment", 3)
add_para(doc, "Phase 1 uses Carla 0.9.13 on the Town04 map in synchronous mode at a fixed 0.05 second timestep at 20 Hz [20]. Town04 was selected for its long straight road segments providing sufficient forward distance for a controlled obstacle approach test. The ego vehicle is spawned at a lane-aligned position with a stationary obstacle 80 metres ahead placed using waypoint alignment. Waypoint-based proportional steering keeps the vehicle tracking the lane centre throughout each 70-second run. Three scenario scripts were developed: dt_phase1_car_front.py for the single obstacle clear baseline, dt_phase1_car_front_rain.py for the extreme rain variant and dt_phase1_car_front_multi_obs_.py for sequential multi-obstacle testing.")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front.py")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front_rain.py")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front_multi_obs_.py")

add_heading(doc, "3.2.2 Simulated Sensors and Fusion", 3)
add_para(doc, "A ray-cast lidar extracts the minimum forward corridor distance using a spatial mask accepting points between 2 m and 120 m forward within 2 m lateral offset. A frontal radar provides per-frame depth and relative velocity for TTC computation. An RGB camera at 1920 by 1080 with 110 degree field of view provides frames for YOLOv8n inference [9]. A gated association method accepts candidate readings only within 12 m of the expected obstacle distance, the fused distance is the minimum across valid candidates and a sentinel value of 999 m suppresses braking when no reading passes the gate [11]. TTC is computed only on valid fused distances as range divided by radar relative velocity [14]. Key threshold constants shared across all Phase 1 scripts are: BRAKE_DISTANCE 8.0 m, AVOID_DISTANCE 15.0 m, TTC_CRITICAL 2.0 seconds and ASSOC_GATE_M 12.0 m.")

add_heading(doc, "3.2.3 Navigation, Safety State Machine and Data Logging", 3)
add_para(doc, "Phase 1 uses lane-following with a planned post-stop avoidance sequence specific to simulation because it depends on the known lane structure in Carla. Six states govern vehicle behaviour:")
add_bullet(doc, "Safe: normal throttle, following the waypoint lane track")
add_bullet(doc, "Warning: fused distance below 15 m, avoidance steering blended with lane following")
add_bullet(doc, "Critical: fused distance below 8 m or TTC below 2 s, throttle set to zero and brake to full")
add_bullet(doc, "Lane Change: left steer at minus 0.25 with throttle 0.35 into the adjacent lane over approximately 60 frames")
add_bullet(doc, "Straightening: blending back to lane steering at throttle 0.45")
add_bullet(doc, "Resumed: normal driving at throttle 0.5 with full waypoint guidance [15]")
add_para(doc, "Each run produces sensor_fusion.csv recording fused distance, individual sensor distances, TTC and safety level per tick; radar_data.csv and lidar_data.csv recording raw sensor outputs; vehicle_state.csv recording speed, throttle, brake, steer and active state; per-frame lidar PLY point clouds; JPEG camera images; and output.mp4 at 20 FPS.")

add_heading(doc, "3.2.4 YOLO Analysis Pipeline", 3)
add_para(doc, "The offline YOLO analysis pipeline loads the recorded RGB camera frames alongside sensor_fusion.csv, vehicle_state.csv and the lidar PLY point clouds. For each frame it performs lane-area estimation, YOLOv8n inference, in-lane filtering of detections, performance logging and output generation. The lane mask is created by converting each frame to HSV, applying colour thresholds targeting road-like surfaces, applying a trapezoid region of interest representing the perspective lane region and using morphological open and close operations to clean the mask. For each detected bounding box the bottom-centre ground contact point is checked against the lane mask and a detection is accepted as in-lane only if lane mask coverage around that point exceeds a set threshold. Eight graph outputs are generated covering inference time over frames, FPS over time, objects detected per frame, top detected object classes, inference time distribution histogram, detections versus inference time scatter, cumulative detections over time and a summary statistics box [9], [22].")
add_para(doc, "A separate visualisation script analyse_logs.py runs offline over the logged CSV files after each scenario and generates seven plots covering sensor distances over time, motor state distribution, fusion level analysis, YOLO camera detections, lidar distance histograms per zone, an avoidance event summary table and overall system performance metrics [21], [22].")
add_github(doc, "Phase1 Carla Project Python Scripts/Visualisation Scripts/Live_Visualisation_with_YOLO.py")
add_github(doc, "Phase1 Carla Project Python Scripts/Visualisation Scripts/analyse_logs.py")

section_banner(doc, "Phase 2: Real Hardware Deployment (Maverick RC Car)", "2E75B6")

add_heading(doc, "3.3 Phase 2 Hardware Components", 2)
add_para(doc, "Phase 2 deploys the validated architecture on a Maverick monster truck chassis with 4WD drive and independent suspension supporting 8 to 12 kg of payload. Table 3.2 lists all hardware components.")

add_table(doc,
    ["Component", "Role", "Key Specification"],
    [
        ["Jetson Orin Nano", "Lidar, camera, fusion, GUI", "8-core Cortex-A78AE, 1024-core Ampere GPU, 8 GB LPDDR5, JetPack 6"],
        ["Raspberry Pi 5", "Radar processing and UDP relay", "4-core Cortex-A76, 8 GB LPDDR4X, DreamHAT+ carrier"],
        ["LD06 Lidar", "360-degree proximity sensing", "230400 baud, 47-byte packets, 12 measurements per packet"],
        ["BGT60TR13C Radar", "60 GHz FMCW moving-target detection", "1 TX, 3 RX, 64 chirps, 128 samples per chirp"],
        ["Arducam IMX477 HQ", "Visual cross-validation with YOLOv8n", "Sony IMX477, CSI-2, 12.3 MP, model B0249"],
        ["Arduino Uno R4 WiFi", "5V PWM motor control bridge", "UDP port 5005, pin 9 ESC, pin 10 steering servo"],
        ["MSC-25RC ESC", "Drive motor throttle actuation", "1700 us forward, 1570 us slow, 1500 us neutral, 1300 us reverse"],
        ["Steering servo", "Front wheel directional control", "1000 us full left, 1500 us centre, 2000 us full right"],
        ["Maverick chassis", "RC vehicle platform", "4WD, independent suspension, 8 to 12 kg payload"],
    ],
    [2.0, 2.2, 2.8]
)
add_caption(doc, "Table 3.2: Phase 2 hardware components, roles and specifications.")

add_para(doc, "The Jetson Orin Nano handles LD06 lidar serial ingestion, YOLOv8n inference on its 1024-core Ampere GPU, sensor fusion and the PyQt5 GUI dashboard. The Raspberry Pi 5 hosts the BGT60TR13C radar on the DreamHAT+ carrier and streams processed track data to the Jetson over UDP on port 9576, with each JSON track record carrying id, range_m, velocity_mps, angle_deg, ttc_s and level [8]. The Arduino receives single-byte UDP commands on port 5005 and outputs 50 Hz PWM to the ESC on pin 9 and the steering servo on pin 10. A key constraint was that the Jetson GPIO outputs 3.3 V which is insufficient for the MSC-25RC ESC requiring 5 V: the Arduino acts as the 5 V signal bridge and the ESC BEC powers the servo rail independently [18].")

add_para(doc, "Figure 3.1 shows the complete RC car platform with all sensors and compute hardware mounted.")
add_image(doc, "fig3_1_rccar_full", "Figure 3.1: RC car with all components mounted and connected.")
add_para(doc, "As shown in Figure 3.1, the chassis provides a stable base with sufficient space for all hardware mounted securely across the upper deck. Figure 3.2 shows the Jetson Orin Nano and LD06 lidar during bench integration before mounting on the vehicle.")
add_image(doc, "fig3_2_jetson_lidar", "Figure 3.2: Jetson Orin Nano and LD06 lidar during bench integration.")
add_para(doc, "As shown in Figure 3.2, the primary compute unit and lidar were validated together on the bench before physical vehicle integration. Figure 3.3 shows the Raspberry Pi 5 with the BGT60TR13C radar on the DreamHAT+ carrier board.")
add_image(doc, "fig3_3_pi_radar", "Figure 3.3: Raspberry Pi 5 with BGT60TR13C radar on DreamHAT+.")
add_para(doc, "As shown in Figure 3.3, the radar sits directly on the DreamHAT+ carrier which interfaces with the Pi 5 via SPI. All radar signal processing runs on the Pi and only clean track data is streamed to the Jetson over the network [8].")
add_github(doc, "Phase2 Hardare Python Scripts/jetson_fusion.py")
add_github(doc, "Phase2 Hardare Python Scripts/radar_send.py")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/arduino_motor_controller/arduino_motor_controller.ino")

add_heading(doc, "3.4 Phase 2 System Power Requirements", 2)
add_para(doc, "The Jetson group typically draws 1.05 to 1.4 A peaking at 2.1 A, consuming 8 to 11 W typically and up to 17 W at peak. The Pi group typically draws 0.95 to 1.15 A peaking at 1.9 A, consuming 4.5 to 5.5 W typically and up to 9 W at peak. The RC car group powered from the 7.2 V LiPo through the ESC BEC peaks above 59 W under full motor load. Total system typical consumption is 18 to 25 W peaking at approximately 85 W. Table 3.3 provides the full power budget.")
add_table(doc,
    ["Component", "Voltage", "Typical Current", "Peak Current", "Typical Power", "Peak Power", "Power Source"],
    [
        ["Jetson Orin Nano", "9 to 20 V", "0.78 to 1.1 A", "1.67 A", "7 to 10 W", "15 W", "DC Barrel"],
        ["Arducam IMX477", "3.3 V", "90 mA", "150 mA", "0.3 W", "0.5 W", "CSI Rail"],
        ["LD06 Lidar", "5 V", "180 mA", "300 mA", "0.9 W", "1.5 W", "USB Port"],
        ["Jetson Group Total", "", "1.05 to 1.4 A", "2.1 A", "8 to 11 W", "17 W", ""],
        ["Raspberry Pi 5", "5 V", "0.8 to 1.0 A", "1.6 A", "4 to 5 W", "8 W", "USB-C"],
        ["BGT60TR13C Radar", "3.3 V", "150 mA", "300 mA", "0.5 W", "1 W", "Pi GPIO Rail"],
        ["Pi Group Total", "", "0.95 to 1.15 A", "1.9 A", "4.5 to 5.5 W", "9 W", ""],
        ["Arduino Uno R4 WiFi", "5 V", "100 mA", "200 mA", "0.5 W", "1 W", "USB / ESC BEC"],
        ["Steering Servo", "5 V", "50 mA", "500 mA", "0.25 W", "2.5 W", "ESC BEC"],
        ["MSC-25RC ESC", "7.2 V", "0.14 to 0.28 A", "3.5 A+", "1 to 2 W", "25 W+", "7.2 V LiPo"],
        ["Drive Motor", "7.2 V", "0.4 to 0.7 A", "4 A+", "3 to 5 W", "30 W+", "7.2 V LiPo"],
        ["Car Group Total", "", "0.7 to 1.1 A", "8.2 A+", "5 to 8 W", "59 W", ""],
        ["System Total", "", "2.7 to 3.7 A", "12.2 A+", "18 to 25 W", "85 W", ""],
    ],
    [1.5, 0.8, 1.1, 1.0, 1.0, 0.9, 1.0]
)
add_caption(doc, "Table 3.3: System power requirements by group.")

add_heading(doc, "3.5 Phase 2 Software Stack", 2)
add_para(doc, "The software stack is distributed across the three compute platforms as follows:")
add_bullet(doc, "Jetson Orin Nano: Ubuntu 22.04, JetPack 6, Python 3.10, PyQt5, PyTorch with CUDA, OpenCV 4.8.0, Ultralytics YOLOv8")
add_bullet(doc, "Raspberry Pi 5: Python 3.11, Infineon ifxAvian SDK providing the BGT60TR13C radar interface utilities")
add_bullet(doc, "Arduino Uno R4 WiFi: C++ firmware using the Arduino Servo library for 50 Hz PWM generation")
add_para(doc, "All inter-device communication uses standard UDP sockets over the shared WiFi hotspot with no additional middleware [22].")

add_heading(doc, "3.6 Phase 2 Distributed Architecture and Threading", 2)
add_para(doc, "The system distributes work across three boards over a shared WiFi hotspot. The Pi runs the radar pipeline and streams confirmed track data to the Jetson on UDP port 9576. The Jetson runs lidar parsing, YOLOv8n inference, sensor fusion, motor decisions and data logging. The Arduino translates UDP commands into ESC and servo PWM. The Arduino Uno R4 WiFi is configured with a static IP address of 172.20.10.3 set using WiFi.config before WiFi.begin, ensuring a consistent address on every power cycle without router configuration [22]. The Jetson also listens on a separate UDP port for pause and resume commands from a laptop, allowing immediate halt during testing.")
add_para(doc, "The fusion application runs four concurrent threads: the lidar thread reads the LD06 serial port and updates shared state under a threading.Lock; the radar thread listens on UDP port 9576 and updates shared radar state as JSON packets arrive; the camera thread captures 416 by 416 frames through GStreamer nvarguscamerasrc with nvvidconv hardware resize and runs YOLOv8n on CUDA every second frame to manage thermal load; the GUI thread polls shared state every 100 ms via QTimer and re-renders the dashboard without waiting on sensor data, sustaining approximately 9.5 FPS [22]. Real test logs confirm the fusion loop running at approximately 15 Hz with consistent 60 to 65 ms timestamps and no UDP lag.")
add_para(doc, "Figure 3.4 shows the fusion GUI dashboard during operation.")
add_image(doc, "fig3_4_fusion_gui", "Figure 3.4: Fusion GUI showing lidar scatter panel, radar range-velocity panel, TTC chart and status bar.")
add_para(doc, "As shown in Figure 3.4, the left panel renders the lidar bird's-eye scatter plot coloured by zone, the centre panel shows confirmed radar tracks on a range-velocity plot coloured by alert level, the right panel contains the TTC bar chart and track table, and the bottom status bar provides a single-line snapshot of FPS, sensor counts, fused verdict and health state.")

add_heading(doc, "3.7 Phase 2 Sensor Fusion and Health Monitoring", 2)
add_para(doc, "The fusion policy is worst-case-across-sensors: Imminent maps to priority 3, Caution to 2 and Safe to 1. The highest priority sensor governs the output and no sensor can reduce the result [11], [13]. Lidar handles stationary obstacle detection. Radar handles moving targets through TTC. The camera cross-validates lidar alerts: if YOLO finds nothing in the central 40 percent of the frame the lidar alert is suppressed as noise, and if both agree the detection is confirmed [12]. The camera cannot trigger alerts independently. The health monitor floors the output at Caution if both lidar and camera are simultaneously impaired ensuring the system never returns Safe when two primary sensors are unavailable [13].")
add_para(doc, "The health layer distinguishes a disconnected sensor from a degraded one. A disconnected sensor stops sending data and its timestamp goes stale. A degraded sensor is still transmitting but cannot detect anything useful: a lidar with a blocked lens continues sending packets with zero valid forward-zone points, or a camera in a scene too dark for detection [3], [4]. Lidar degradation is flagged after 15 consecutive frames of zero forward-zone points and triggers radar static fallback coverage. Camera degradation triggers when mean brightness falls below 40 and current detections are cleared immediately.")

add_heading(doc, "3.8 Phase 2 LD06 Lidar Pipeline", 2)
add_para(doc, "The LD06 transmits 47-byte packets at 230400 baud with 12 distance measurements per packet. Bytes 4 to 5 encode the start angle and bytes 42 to 43 encode the end angle, both divided by 100 to convert to degrees, with each measurement occupying 3 bytes for distance low byte, distance high byte and confidence. When first connected every obstacle appeared behind the origin because the sensor native zero-degree reference points opposite to the vehicle forward direction. An angle offset correction applied as (angle + ANGLE_OFFSET_DEG) % 360.0 and a Y-axis sign inversion of minus 1 resolved this [1], [2]. Three filtering layers were then applied in sequence:")
add_bullet(doc, "FOV restriction to plus or minus 20 degrees forward, a 20 cm minimum distance floor and a secondary geometric check for near-origin stray points [2]")
add_bullet(doc, "Minimum point count per zone scaled with detection distance so closer objects need more points to register than distant ones [3]")
add_bullet(doc, "Persistence filter requiring 3 consecutive scans to raise an alert and 4 consecutive clear scans to lower it, with asymmetric hysteresis preventing state flickering [4]")
add_para(doc, "These three layers removed virtually all false positives observed during bench testing. Two additional threshold changes were applied based on evidence from field test logs. The centre zone width was reduced from 19 cm to 12 cm total because a bottle is approximately 8 cm wide and the wider zone caused thin objects placed slightly off-centre to register in both the centre and left zones simultaneously. The stop threshold was raised from 0.35 m to 0.6 m to match real sensor readings: the lidar was consistently reading 0.44 to 0.48 m when physically blocked at close range and with the original 0.35 m threshold the reverse manoeuvre was never triggering [22]. Table 3.4 defines the three detection zones.")
add_table(doc,
    ["Zone", "Angular Range", "Navigation Response"],
    [
        ["Centre (red)", "340 to 360 and 0 to 20 degrees", "Reduce speed or stop"],
        ["Left (blue)", "20 to 90 degrees", "Steer right to avoid"],
        ["Right (green)", "270 to 340 degrees", "Steer left to avoid"],
    ],
    [2.0, 3.0, 2.0]
)
add_caption(doc, "Table 3.4: Lidar zone definitions and navigation responses.")
add_para(doc, "Figure 3.5 shows the corrected live scatter plot with points coloured by zone and threshold lines overlaid.")
add_image(doc, "fig3_5_lidar_scatter", "Figure 3.5: LD06 live scatter plot with three-zone colouring and threshold lines.")
add_para(doc, "As shown in Figure 3.5, the centre zone in red, left in blue and right in green are clearly distinguishable in the bird's-eye view with the two horizontal threshold lines marking the stop and slow distances providing a real-time reference for when detections are entering the danger range.")
add_github(doc, "Phase2 Hardare Python Scripts/initial_Test_Debug_Scripts/lidar_GUI.py")

add_heading(doc, "3.9 Phase 2 BGT60TR13C Radar Pipeline", 2)
add_para(doc, "The BGT60TR13C runs on the Pi 5 with one transmit and three receive antennas producing a raw data array of shape (3, 64, 128) per frame. Five processing layers are applied in sequence [5], [7], [8]:")
add_bullet(doc, "Three-stage FFT: Range FFT along the sample axis gives distance, Doppler FFT along the chirp axis gives velocity and Angle FFT across the three receivers gives bearing")
add_bullet(doc, "MTI background subtraction at memory coefficient 0.92 removes static clutter from walls, floors and the chassis, requiring approximately 15 frames to settle [5]")
add_bullet(doc, "CA-CFAR detection at 16 dB identifies candidates above the adaptive local noise floor [6]")
add_bullet(doc, "Nearest-neighbour tracking confirms tracks after 3 consecutive matched frames and deletes after 3 consecutive misses")
add_bullet(doc, "Approach confirmation gate requiring 3 consecutive frames of consistently negative radial velocity before a track triggers a warning [7]")
add_para(doc, "TTC triggers Imminent below 1.5 seconds or within 0.5 m and Caution below 3.0 seconds or within 1.2 m, with range fallbacks for near-zero velocity [14], [15]. Stationary obstacles are the lidar's responsibility entirely because MTI suppresses static returns by design.")
add_para(doc, "Figure 3.6 shows the radar GUI with a confirmed approaching track displayed on the range-velocity plot.")
add_image(doc, "fig3_6_radar_gui", "Figure 3.6: Radar GUI showing confirmed approaching track with range-velocity position and TTC.")
add_para(doc, "As shown in Figure 3.6, confirmed tracks appear as dots at their range and velocity coordinates coloured by alert level with TTC values annotated alongside each track, allowing the operator to see at a glance which targets are approaching and how urgently.")
add_github(doc, "Phase2 Hardare Python Scripts/radar_send.py")

add_heading(doc, "3.10 Phase 2 Camera Pipeline, Motor Control and Avoidance State Machine", 2)
add_para(doc, "The Arducam IMX477 connects via CSI-2 through a GStreamer nvarguscamerasrc pipeline hardware-resizing to 416 by 416 via nvvidconv. YOLOv8n runs on CUDA every second frame. Only detections whose bounding box centre x falls within the middle 40 percent of the frame are considered. Distance estimation uses projective geometry: estimated distance equals real-world object height multiplied by focal length divided by bounding box pixel height, producing an Imminent, Caution or Safe band without a depth sensor [10]. Camera verdicts are cleared when mean brightness falls below 40. A significant finding from real-world testing was that the IMX477 ribbon cable hanging loose on the front-left side of the chassis entered the lidar scan plane, producing a permanent false left-zone reading of 0.33 to 0.40 m in open space and triggering continuous Imminent state. The fix was to route the cable behind the lidar and raise the sensor on standoffs [22].")
add_para(doc, "Figure 3.7 shows the Arducam IMX477 connected to the Jetson via the CSI-2 port.")
add_image(doc, "fig3_7_camera", "Figure 3.7: Arducam IMX477 connected to Jetson via CSI-2.")
add_para(doc, "As shown in Figure 3.7, the camera connects via ribbon cable directly to the CSI-2 port on the Jetson Orin Nano. The routing of this cable away from the lidar scan plane was identified as a critical hardware fix following real-world testing.")
add_github(doc, "Phase2 Hardare Python Scripts/initial_Test_Debug_Scripts/yolo_test.py")
add_para(doc, "Motor commands are encoded as JSON and sent from the Jetson to the Arduino on UDP port 5005. Safe commands 1700 microseconds forward, Caution commands 1570 microseconds slow forward and Imminent commands 1500 microseconds neutral stop. Steering derives from lidar zone output: a left zone obstacle triggers right steering, a right zone obstacle triggers left steering and both zones clear returns the servo to 1500 microseconds centre.")
add_para(doc, "The motor controller implements a six-state avoidance machine running at 20 Hz. The key design principle is sensor-driven state transitions rather than fixed timers: states exit as soon as sensor conditions are met rather than running for a predetermined duration. Reversing is the only state retaining a fixed duration of 2.5 seconds to guarantee physical clearance before a turn is attempted. Normal operates across three distance zones: above 1.5 m the car drives at full forward speed, between 1.0 m and 1.5 m it slows and steers gently away from the obstacle, and between 0.6 m and 1.0 m it steers around the obstacle at slow speed without stopping. Only when an obstacle is under 0.6 m and both sides are simultaneously blocked does the car stop and reverse. Stopping holds full brake and selects an escape direction. Reversing drives straight reverse for 2.5 seconds. Turning applies slow forward toward the escape direction and exits the moment the front centre distance reads above 1.0 m. Escaping drives slow forward on the new heading, replanning if a new obstacle appears after a 0.5 second grace period, and straightening temporarily if a side wall falls within 0.25 m. Straightening recentres steering and resets retry counters. A stuck detection timer handles scenarios where the car has been continuously steering around an obstacle for more than 4 seconds without clearing the front [22].")
add_para(doc, "Figure 3.8 shows the updated sensor-driven avoidance state machine diagram.")
add_image(doc, "fig3_8_statemachine", "Figure 3.8: Sensor-driven avoidance state machine with six states, sensor-driven exits and retry logic.")
add_para(doc, "As shown in Figure 3.8, the state machine shows all six states with their timing constraints, sensor-driven exit conditions, and the retry and direction flip paths that activate when the chosen escape direction is repeatedly blocked.")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/motor_controller.py")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/car_control.py")

add_heading(doc, "3.11 Phase 2 ESC Reverse Arming and Autostart", 2)
add_para(doc, "The MSC-25RC ESC requires a specific three-step arming sequence before reverse engages: a brake pulse at 1300 microseconds for 200 ms, a neutral pulse at 1500 microseconds for 300 ms, then the reverse pulse at 1300 microseconds, taking approximately 500 ms total. The motor controller was sending the reverse command every 50 ms at 20 Hz, restarting the arming sequence each time so the car never physically reversed. A reverseArmed flag on the Arduino fixed this: once the sequence completes all subsequent reverse commands are ignored and the ESC holds reverse PWM continuously until a different command arrives [22].")
add_para(doc, "Figure 3.9 shows the ESC reverse arming sequence and the reverseArmed flag logic.")
add_image(doc, "fig3_9_esc_arming", "Figure 3.9: ESC reverse arming sequence showing IDLE to BRAKE to NEUTRAL WAIT to REVERSING with reverseArmed flag.")
add_para(doc, "As shown in Figure 3.9, without the reverseArmed flag the arming sequence restarts on every incoming command and the ESC never reaches the Reversing state. With the flag in place the sequence completes once and reverse is maintained until a different motor command arrives.")
add_para(doc, "Both boards start automatically on power-on: the radar streamer on the Pi runs as a systemd service and the Jetson launches the fusion GUI via a desktop autostart file, requiring no laptop for field operation.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "Chapter 4: Performance Evaluation Setup and Scenarios", 1)

add_heading(doc, "4.1 Evaluation Methodology", 2)
add_para(doc, "The test structure follows a progressive validation methodology. Phase 1 establishes a simulation baseline through three scenarios of increasing complexity: a clear conditions single-obstacle test to validate the core pipeline under ideal conditions, a rain variant holding all other variables constant to isolate sensor redundancy by stressing the camera, and a multi-obstacle sequential test to verify state machine integrity over repeated encounters. This mirrors controlled ablation study design where one variable changes at a time to attribute results to specific system components [11], [13]. Phase 2 bench testing preceded on-car testing as a safety precondition to localise faults to the integration layer when they appear. The two on-car sessions are structured as an extended multi-obstacle stress test and a single-obstacle controlled reproduction of the Phase 1 baseline, allowing direct comparison between simulation and hardware performance.")
add_para(doc, "Metrics assessed across both phases are detection rate as true positives over total obstacles encountered, false positive rate as spurious alerts in clear conditions, TTC accuracy as estimated versus measured approach time, reaction latency from fusion alert to first motor command, stop distance at halt, steering accuracy, drive time efficiency as proportion of session time in normal forward driving and avoidance success rate as the proportion of encounters completing the full manoeuvre sequence. Phase 1 metrics come from the CSV log files and Phase 2 metrics from the on-board data logger.")

section_banner(doc, "Phase 1 Evaluation: Carla Simulation Scenarios", "1A6E3C")

add_heading(doc, "4.2 Phase 1 Test Environment", 2)
add_para(doc, "All Phase 1 runs execute in Carla on Town04 in synchronous mode at 20 Hz, reset to a clean state between runs. Ego vehicle spawn, obstacle placement and threshold constants are identical across all runs so the only variable changing between scenarios is weather conditions and obstacle count [20], [21]. In Phase 1 each scenario script constitutes both the test configuration and the test execution. Running the script under the defined conditions produces the logged CSV outputs from which all Phase 1 metrics are derived. This differs from Phase 2 where the test environment is physical and results reflect real-world sensor behaviour rather than simulated sensor models.")

add_heading(doc, "Scenario 1: Single Obstacle, Clear Conditions", 2)
add_heading(doc, "4.3 Scenario 1 Setup", 3)
add_para(doc, "A single stationary obstacle vehicle is placed 80 metres ahead in the ego lane under default clear weather with all three simulated sensors operating normally. This is the clean pipeline baseline against which all other scenario results are compared. The ego vehicle approaches from the spawn point at normal throttle following the waypoint lane track. As the fused distance drops below 15 m the system enters Warning and blends avoidance steering with lane following. When fused distance drops below 8 m or TTC drops below 2 seconds the system transitions to Critical and applies full emergency braking. After the vehicle stops, Lane Change is initiated with left steer at minus 0.25 and throttle 0.35. Straightening then blends back to lane steering at throttle 0.45 and Resumed restores normal driving at throttle 0.5 [20], [22].")
add_para(doc, "Figure 4.1 shows the Town04 initial setup with the ego vehicle and stationary obstacle aligned in the driving lane.")
add_image(doc, "fig4_1_town04", "Figure 4.1: Town04 initial setup with ego vehicle and obstacle aligned in lane with sensor overlays.")
add_para(doc, "As shown in Figure 4.1, the simulation provides a clear straight approach path with the obstacle placed directly ahead in the same lane, giving the ego vehicle sufficient forward distance to detect and respond before reaching the critical distance threshold.")
add_para(doc, "Figure 4.2 shows the emergency brake trigger moment under clear conditions.")
add_image(doc, "fig4_2_brake_clear", "Figure 4.2: Emergency brake trigger moment under clear conditions.")
add_para(doc, "As shown in Figure 4.2, the vehicle has transitioned to Critical state with the fused distance having dropped below 8 m. Throttle is set to zero and full braking is applied, bringing the vehicle to a complete stop before contact with the obstacle.")
add_para(doc, "Figure 4.3 shows the lane change manoeuvre beginning after the vehicle has stopped.")
add_image(doc, "fig4_3_lane_change", "Figure 4.3: Lane change manoeuvre, Scenario 1.")
add_para(doc, "As shown in Figure 4.3, the vehicle applies gentle left steering at minus 0.25 with throttle 0.35 to begin moving into the adjacent lane, clearing the obstacle without reversing.")
add_para(doc, "Figure 4.4 shows the vehicle straightening after completing the lane change.")
add_image(doc, "fig4_4_straightening", "Figure 4.4: Vehicle straightening after lane change.")
add_para(doc, "As shown in Figure 4.4, the steering is blending back toward the lane-following signal at throttle 0.45 as the vehicle realigns with the new lane ahead of the obstacle.")
add_para(doc, "Figure 4.5 shows the vehicle resuming normal lane driving.")
add_image(doc, "fig4_5_resume", "Figure 4.5: Vehicle resuming normal lane driving.")
add_para(doc, "As shown in Figure 4.5, the vehicle has returned to the Resumed state with full throttle at 0.5 and complete waypoint-guided steering restored, driving normally past the cleared obstacle.")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front.py")

add_heading(doc, "4.4 Scenario 1 Results and Analysis", 3)
add_para(doc, "Radar readings fluctuated around 12 to 14 m during the approach, reflecting characteristics of the Carla radar simulation model. Lidar was considerably more stable and settled at approximately 4.3 m at closest range. The fused distance correctly overrode the noisier radar readings through the 12 m association gate throughout the approach [11]. Both the distance and TTC thresholds converged at the same braking event, confirming the dual-trigger logic is internally consistent and neither condition is redundant. Safety state distribution across the 70-second run was approximately 89 percent Safe, 2 to 3 percent Warning and 8 percent Critical. The radar to lidar mean distance difference of approximately 8 to 9 m was absorbed correctly by the association gate. TTC values remained above 2 seconds throughout the run with the sub-threshold cluster appearing exclusively at the braking event. Lane change, straightening and resume completed without collision. Scenario 1 is a pass.")
add_para(doc, "Figure 4.6 shows the performance dashboard generated from the Scenario 1 log files.")
add_image(doc, "fig4_6_dashboard", "Figure 4.6: Scenario 1 performance dashboard showing sensor distances over time, TTC profile and safety state distribution.")
add_para(doc, "As shown in Figure 4.6, the sensor distances over time graph confirms radar fluctuating above lidar during the approach with the fused distance correctly tracking the lower validated reading. The TTC graph shows a sharp drop through the 2 second threshold at the braking event. The safety state pie chart confirms the dominant Safe proportion at approximately 89 percent with the small Critical region corresponding to the braking and post-stop period. The radar to lidar distance agreement histogram confirms the mean difference of approximately 8 to 9 m which the association gate absorbed correctly throughout the run.")

add_heading(doc, "Scenario 2: Single Obstacle, Extreme Rain", 2)
add_heading(doc, "4.5 Scenario 2 Setup", 3)
add_para(doc, "Scenario 2 repeats the identical approach, braking, lane change and resume sequence as Scenario 1 with the only variable being extreme adverse weather applied through Carla's WeatherParameters API. Precipitation is set to 100 percent, road wetness to 100 percent, fog density to 40 percent with onset at 30 m and sun altitude to minus 30 degrees below the horizon. Headlights are enabled automatically. The purpose of Scenario 2 is to isolate sensor redundancy: if the braking and recovery sequence matches Scenario 1 despite heavy camera degradation it proves that lidar and radar alone are driving the safety decision independently of camera state [20], [22].")
add_table(doc,
    ["Parameter", "Scenario 1 Clear", "Scenario 2 Extreme Rain"],
    [
        ["Precipitation", "None", "100 percent"],
        ["Precipitation deposits", "None", "100 percent, wet roads"],
        ["Sun altitude angle", "Default", "Minus 30 degrees, below horizon"],
        ["Fog density", "None", "40 percent"],
        ["Fog start distance", "Not applicable", "30 m"],
        ["Road wetness", "Dry", "100 percent"],
        ["Headlights", "Off", "On"],
    ],
    [2.5, 2.0, 2.5]
)
add_caption(doc, "Table 4.1: Weather parameter comparison between Scenario 1 clear and Scenario 2 extreme rain.")
add_para(doc, "Figure 4.7 shows the extreme rain simulation environment.")
add_image(doc, "fig4_7_rain_setup", "Figure 4.7: Extreme rain setup showing fog, wet roads and reduced visibility.")
add_para(doc, "As shown in Figure 4.7, the combination of fog starting at 30 m and the sun altitude set to minus 30 degrees below the horizon creates substantially reduced forward camera visibility. The wet road surface and precipitation overlay are clearly visible across the scene.")
add_para(doc, "Figure 4.8 shows the braking event under extreme rain conditions.")
add_image(doc, "fig4_8_brake_rain", "Figure 4.8: Braking under extreme rain conditions.")
add_para(doc, "As shown in Figure 4.8, despite the heavily degraded camera stream the system correctly transitions to Critical state and applies full emergency braking at the same fused distance threshold as Scenario 1, driven by lidar and radar independently of camera state.")
add_para(doc, "Figure 4.9 shows the lane change and resume under extreme rain.")
add_image(doc, "fig4_9_lanechange_rain", "Figure 4.9: Lane change and resume under extreme rain.")
add_para(doc, "As shown in Figure 4.9, the lane change, straightening and resume sequence completes correctly under rain conditions with no difference in behaviour compared to the clear conditions baseline, confirming the planned post-stop recovery is unaffected by weather.")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front_rain.py")

add_heading(doc, "4.6 Scenario 2 Results and Analysis", 3)
add_para(doc, "The camera stream was heavily degraded throughout the run. Fog onset at 30 m cut forward visible range substantially and the precipitation overlay reduced scene contrast across the frame. Lidar distance extraction remained stable throughout because the ray-cast lidar model in Carla does not simulate precipitation attenuation, consistent with real lidar behaviour in light to moderate rain. Radar readings were equally unaffected by visibility conditions [8]. The fusion pipeline produced the same Warning and Critical transitions at the same threshold values as Scenario 1 with no meaningful delay. This demonstrates directly that the multi-sensor architecture maintains safety-critical detection under conditions where a camera-only system would struggle substantially at range. The lane change and recovery sequence completed correctly under rain. Scenario 2 is a pass.")

add_heading(doc, "Scenario 3: Sequential Multi-Obstacle", 2)
add_heading(doc, "4.7 Scenario 3 Setup", 3)
add_para(doc, "Scenario 3 tests whether the safety state machine resets correctly between successive avoidance cycles without state contamination from previous encounters. Three obstacle vehicles are configured with the first placed at 150 m ahead in the driving lane. Each subsequent obstacle is spawned dynamically 150 m ahead in the current driving lane as the preceding one is passed. Obstacles are assigned distinct colours: red, green and blue respectively for visual identification in the output video. An obstacle is classified as passed when the forward dot product of the ego-to-obstacle vector becomes negative or zero, at which point the braked and lane_change_complete flags reset and the next obstacle is treated as a completely fresh event [20], [22].")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front_multi_obs_.py")

add_heading(doc, "4.8 Scenario 3 Results and Analysis", 3)
add_para(doc, "All three obstacles spawned correctly in the driving lane at 150 m ahead. Each triggered a fully independent detection, Warning transition, Critical transition, emergency brake and lane change sequence with no evidence of state contamination from the previous cycle. The dot product classification correctly identified each obstacle as passed once cleared and the flag reset allowed the state machine to respond to each subsequent obstacle as a fresh event. Between cycles the vehicle returned cleanly to lane-following at throttle 0.5 with no residual steering anomaly. The output video Multi Object Avoidance.mp4 confirms all three avoidance cycles visually. Scenario 3 is a pass.")

add_heading(doc, "4.9 Phase 1 YOLO Performance Evaluation", 2)
add_heading(doc, "4.9.1 YOLOv8n Baseline Evaluation", 3)
add_para(doc, "YOLOv8n was evaluated offline over the recorded camera frames from all Phase 1 scenarios using Live_Visualisation_with_YOLO.py. Mean inference time was approximately 9.2 ms per frame corresponding to approximately 114 FPS after GPU warm-up, with an initial spike of approximately 900 ms at the first frame for model loading and GPU memory allocation. Average in-lane detection count was approximately 0.12 per frame with car as the dominant class, consistent with the Town04 obstacle type. These results confirm YOLO is not a latency bottleneck in the Phase 1 pipeline [9].")

add_heading(doc, "4.9.2 YOLOv5n versus YOLOv8n Comparative Evaluation", 3)
add_para(doc, "As an extension to the baseline YOLO evaluation, YOLOv5n and YOLOv8n were compared across both clear and rain scenario conditions to determine which model provides superior real-time performance. Both models used the identical lane-relevance filtering pipeline across 2,331 frames for the clear conditions scenario and 1,185 frames for the rain scenario [9], [22].")
add_para(doc, "Under clear conditions YOLOv8n achieved a mean inference time of 19.8 ms against 20.5 ms for YOLOv5n, corresponding to mean FPS of 50.6 versus 49.4. YOLOv5n produced a catastrophic spike of 673.90 ms during the clear conditions run, representing a 33 times slowdown from its typical performance. At 10 m/s vehicle speed this creates an estimated 6.7 m blind spot. YOLOv8n's worst-case inference time was 67.88 ms, 10 times more stable, creating only a 0.68 m blind spot under equivalent conditions. Standard deviation of inference time was 13.56 ms for YOLOv5n against 1.18 ms for YOLOv8n. Car detections were 133 for YOLOv8n versus 129 for YOLOv5n [9].")
add_para(doc, "Under extreme rain conditions YOLOv5n detected 28 cars versus 21 for YOLOv8n. YOLOv8n produced 7 unique detected classes versus 5 for YOLOv5n, with the higher class count reflecting more misclassifications under degraded frames. Table 4.2 summarises the full comparison.")
add_table(doc,
    ["Metric", "YOLOv5n Clear", "YOLOv8n Clear", "YOLOv5n Rain", "YOLOv8n Rain"],
    [
        ["Car detections", "129", "133", "28", "21"],
        ["Mean inference time", "20.5 ms", "19.8 ms", "19.9 ms", "19.7 ms"],
        ["Mean FPS", "49.4", "50.6", "50.3", "51.0"],
        ["Min FPS", "1.5", "10.9", "14.5", "17.7"],
        ["Max inference time", "673.90 ms", "67.88 ms", "68.93 ms", "56.49 ms"],
        ["Std deviation", "13.56 ms", "1.18 ms", "1.55 ms", "not stated"],
        ["Unique classes", "1", "1", "5", "7"],
    ],
    [2.0, 1.4, 1.4, 1.4, 1.4]
)
add_caption(doc, "Table 4.2: YOLOv5n versus YOLOv8n performance comparison under clear and rain conditions.")
add_para(doc, "YOLOv8n was selected for the final pipeline on the basis of its significantly superior inference stability under clear conditions. The catastrophic single-frame spike produced by YOLOv5n is unacceptable for a real-time collision avoidance system [9]. Figures 4.10 to 4.17 show all eight performance and statistical analysis dashboards for both models across both conditions.")

add_para(doc, "Figure 4.10 shows the YOLOv8n baseline performance dashboard from the initial evaluation.")
add_image(doc, "yolo_v8_baseline", "Figure 4.10: YOLOv8n baseline performance dashboard showing inference time per frame, FPS over time, detections per frame and top detected classes.")
add_para(doc, "As shown in Figure 4.10, the YOLOv8n baseline evaluation confirms mean inference time of approximately 9.2 ms per frame corresponding to approximately 114 FPS after GPU warm-up with an initial spike for model loading. Average in-lane detection count is approximately 0.12 per frame with car as the dominant class.")

add_para(doc, "Figure 4.11 shows the YOLOv5n clear conditions performance graphs.")
add_image(doc, "yolo_v5_clear_perf", "Figure 4.11: YOLOv5n clear conditions performance graphs showing inference time per frame, FPS over time, objects detected per frame and top detected classes.")
add_para(doc, "As shown in Figure 4.11, the YOLOv5n inference time graph under clear conditions shows a mostly stable baseline around 20 ms with one catastrophic spike reaching 673 ms, representing a 33 times slowdown. The FPS graph drops to a minimum of 1.5 FPS at that spike. The class distribution confirms car as the only detected class with 129 total detections.")

add_para(doc, "Figure 4.12 shows the YOLOv5n clear conditions statistical analysis.")
add_image(doc, "yolo_v5_clear_stats", "Figure 4.12: YOLOv5n clear conditions statistical analysis showing inference time distribution histogram, detections vs processing time scatter plot, cumulative detections and performance summary box.")
add_para(doc, "As shown in Figure 4.12, the inference time distribution is highly concentrated around 20 ms with a small secondary peak at 23 ms. The detections versus processing time scatter plot shows a nearly flat trend line of y equals 0.15x plus 20.52, confirming the 673 ms spike occurs independently of detection complexity. Cumulative detections reach 129 cars. The performance summary confirms mean inference of 20.52 ms, standard deviation of 13.56 ms, maximum of 673.90 ms and detection rate of 5.5 percent.")

add_para(doc, "Figure 4.13 shows the YOLOv8n clear conditions performance graphs.")
add_image(doc, "yolo_v8_clear_perf", "Figure 4.13: YOLOv8n clear conditions performance graphs showing inference time per frame, FPS over time, objects detected per frame and top detected classes.")
add_para(doc, "As shown in Figure 4.13, the YOLOv8n inference time graph under clear conditions shows an extremely stable baseline throughout the entire run with no catastrophic spikes. The FPS minimum is 10.9 compared to YOLOv5n at 1.5. The detection pattern shows 133 total car detections, 4 more than YOLOv5n under identical conditions.")

add_para(doc, "Figure 4.14 shows the YOLOv8n clear conditions statistical analysis.")
add_image(doc, "yolo_v8_clear_stats", "Figure 4.14: YOLOv8n clear conditions statistical analysis showing inference time distribution histogram, detections vs processing time scatter plot, cumulative detections and performance summary box.")
add_para(doc, "As shown in Figure 4.14, the inference time distribution is very tight and concentrated on 20 ms with no secondary peaks or outliers. The trend line of y equals 0.44x plus 19.92 shows no meaningful correlation between detection count and processing time. The performance summary confirms mean inference of 19.95 ms, standard deviation of only 1.18 ms compared to YOLOv5n at 13.56 ms, and a maximum of 67.88 ms versus 673.90 ms, representing a 10 times improvement in worst-case stability.")

add_para(doc, "Figure 4.15 shows the YOLOv5n rain conditions performance graphs.")
add_image(doc, "yolo_v5_rain_perf", "Figure 4.15: YOLOv5n rain conditions performance graphs showing inference time per frame, FPS over time, objects detected per frame and top detected classes.")
add_para(doc, "As shown in Figure 4.15, under rain conditions YOLOv5n shows a stable baseline around 20 ms with a peak spike reaching only 69 ms, dramatically better than the 673 ms spike in clear conditions. Mean FPS is 50.3 with a minimum of 14.5. The detected classes show 28 cars alongside airplane at 19, train at 11 and a few other misclassifications, reflecting the reduced lane filter effectiveness under degraded image quality.")

add_para(doc, "Figure 4.16 shows the YOLOv5n rain conditions statistical analysis.")
add_image(doc, "yolo_v5_rain_stats", "Figure 4.16: YOLOv5n rain conditions statistical analysis showing inference time distribution histogram, detections vs processing time scatter plot, cumulative detections and performance summary box.")
add_para(doc, "As shown in Figure 4.16, the inference time distribution under rain is tight around 20 ms with minimal spread and no significant outliers. The scatter plot trend line of y equals 0.36x plus 19.92 shows no correlation between object count and inference time. Cumulative detections reach 61 total objects of which 28 are cars. The performance summary confirms mean inference of 19.94 ms, standard deviation of 1.55 ms and a maximum of 68.93 ms.")

add_para(doc, "Figure 4.17 shows the YOLOv8n rain conditions performance graphs.")
add_image(doc, "yolo_v8_rain_perf", "Figure 4.17: YOLOv8n rain conditions performance graphs showing inference time per frame, FPS over time, objects detected per frame and top detected classes.")
add_para(doc, "As shown in Figure 4.17, under rain conditions YOLOv8n maintains a stable 20 ms baseline with a peak around 56 ms. Mean FPS is 51.0 with a minimum of 17.7, better than YOLOv5n at 14.5. However the detected classes reveal the key weakness: car accounts for only 21 detections while airplane reaches 49 and baseball bat 24, with 7 unique classes total versus 5 for YOLOv5n.")

add_para(doc, "Figure 4.18 shows the YOLOv8n rain conditions statistical analysis.")
add_image(doc, "yolo_v8_rain_stats", "Figure 4.18: YOLOv8n rain conditions statistical analysis showing inference time distribution histogram, detections vs processing time scatter plot, cumulative detections and performance summary box.")
add_para(doc, "As shown in Figure 4.18, the inference time distribution under rain shows a sharp peak at 19 to 20 ms consistent with the stable clear conditions performance. Cumulative detections reach 104 total objects across all classes compared to YOLOv5n at 61. The high total object count reflects 49 airplane and 24 baseball bat misclassifications dominating the class distribution. This is the critical finding from the rain evaluation: while YOLOv8n is superior in stable clear conditions, its anchor-free design is more susceptible to misclassification under visual degradation, making the cross-validation role of the camera even more important in the fusion pipeline under adverse weather.")

add_github(doc, "Phase1 Carla Project Python Scripts/Visualisation Scripts/Live_Visualisation_with_YOLO.py")
add_github(doc, "Phase1 Carla Project Python Scripts/Visualisation Scripts/analyse_logs.py")

section_banner(doc, "Phase 2: Real Hardware Deployment (Maverick RC Car)", "2E75B6")

add_heading(doc, "3.3 Phase 2 Hardware Components", 2)
add_para(doc, "Phase 2 deploys the validated architecture on a Maverick monster truck chassis with 4WD drive and independent suspension supporting 8 to 12 kg of payload. Table 3.2 lists all hardware components.")

add_table(doc,
    ["Component", "Role", "Key Specification"],
    [
        ["Jetson Orin Nano", "Lidar, camera, fusion, GUI", "8-core Cortex-A78AE, 1024-core Ampere GPU, 8 GB LPDDR5, JetPack 6"],
        ["Raspberry Pi 5", "Radar processing and UDP relay", "4-core Cortex-A76, 8 GB LPDDR4X, DreamHAT+ carrier"],
        ["LD06 Lidar", "360-degree proximity sensing", "230400 baud, 47-byte packets, 12 measurements per packet"],
        ["BGT60TR13C Radar", "60 GHz FMCW moving-target detection", "1 TX, 3 RX, 64 chirps, 128 samples per chirp"],
        ["Arducam IMX477 HQ", "Visual cross-validation with YOLOv8n", "Sony IMX477, CSI-2, 12.3 MP, model B0249"],
        ["Arduino Uno R4 WiFi", "5V PWM motor control bridge", "UDP port 5005, pin 9 ESC, pin 10 steering servo"],
        ["MSC-25RC ESC", "Drive motor throttle actuation", "1700 us forward, 1570 us slow, 1500 us neutral, 1300 us reverse"],
        ["Steering servo", "Front wheel directional control", "1000 us full left, 1500 us centre, 2000 us full right"],
        ["Maverick chassis", "RC vehicle platform", "4WD, independent suspension, 8 to 12 kg payload"],
    ],
    [2.0, 2.2, 2.8]
)
add_caption(doc, "Table 3.2: Phase 2 hardware components, roles and specifications.")

add_para(doc, "The Jetson Orin Nano handles LD06 lidar serial ingestion, YOLOv8n inference on its 1024-core Ampere GPU, sensor fusion and the PyQt5 GUI dashboard. The Raspberry Pi 5 hosts the BGT60TR13C radar on the DreamHAT+ carrier and streams processed track data to the Jetson over UDP on port 9576, with each JSON track record carrying id, range_m, velocity_mps, angle_deg, ttc_s and level [8]. The Arduino receives single-byte UDP commands on port 5005 and outputs 50 Hz PWM to the ESC on pin 9 and the steering servo on pin 10. A key constraint was that the Jetson GPIO outputs 3.3 V which is insufficient for the MSC-25RC ESC requiring 5 V: the Arduino acts as the 5 V signal bridge and the ESC BEC powers the servo rail independently [18].")

add_para(doc, "Figure 3.1 shows the complete RC car platform with all sensors and compute hardware mounted.")
add_image(doc, "fig3_1_rccar_full", "Figure 3.1: RC car with all components mounted and connected.")
add_para(doc, "As shown in Figure 3.1, the chassis provides a stable base with sufficient space for all hardware mounted securely across the upper deck. Figure 3.2 shows the Jetson Orin Nano and LD06 lidar during bench integration before mounting on the vehicle.")
add_image(doc, "fig3_2_jetson_lidar", "Figure 3.2: Jetson Orin Nano and LD06 lidar during bench integration.")
add_para(doc, "As shown in Figure 3.2, the primary compute unit and lidar were validated together on the bench before physical vehicle integration. Figure 3.3 shows the Raspberry Pi 5 with the BGT60TR13C radar on the DreamHAT+ carrier board.")
add_image(doc, "fig3_3_pi_radar", "Figure 3.3: Raspberry Pi 5 with BGT60TR13C radar on DreamHAT+.")
add_para(doc, "As shown in Figure 3.3, the radar sits directly on the DreamHAT+ carrier which interfaces with the Pi 5 via SPI. All radar signal processing runs on the Pi and only clean track data is streamed to the Jetson over the network [8].")
add_github(doc, "Phase2 Hardare Python Scripts/jetson_fusion.py")
add_github(doc, "Phase2 Hardare Python Scripts/radar_send.py")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/arduino_motor_controller/arduino_motor_controller.ino")

add_heading(doc, "3.4 Phase 2 System Power Requirements", 2)
add_para(doc, "The Jetson group typically draws 1.05 to 1.4 A peaking at 2.1 A, consuming 8 to 11 W typically and up to 17 W at peak. The Pi group typically draws 0.95 to 1.15 A peaking at 1.9 A, consuming 4.5 to 5.5 W typically and up to 9 W at peak. The RC car group powered from the 7.2 V LiPo through the ESC BEC peaks above 59 W under full motor load. Total system typical consumption is 18 to 25 W peaking at approximately 85 W. Table 3.3 provides the full power budget.")
add_table(doc,
    ["Component", "Voltage", "Typical Current", "Peak Current", "Typical Power", "Peak Power", "Power Source"],
    [
        ["Jetson Orin Nano", "9 to 20 V", "0.78 to 1.1 A", "1.67 A", "7 to 10 W", "15 W", "DC Barrel"],
        ["Arducam IMX477", "3.3 V", "90 mA", "150 mA", "0.3 W", "0.5 W", "CSI Rail"],
        ["LD06 Lidar", "5 V", "180 mA", "300 mA", "0.9 W", "1.5 W", "USB Port"],
        ["Jetson Group Total", "", "1.05 to 1.4 A", "2.1 A", "8 to 11 W", "17 W", ""],
        ["Raspberry Pi 5", "5 V", "0.8 to 1.0 A", "1.6 A", "4 to 5 W", "8 W", "USB-C"],
        ["BGT60TR13C Radar", "3.3 V", "150 mA", "300 mA", "0.5 W", "1 W", "Pi GPIO Rail"],
        ["Pi Group Total", "", "0.95 to 1.15 A", "1.9 A", "4.5 to 5.5 W", "9 W", ""],
        ["Arduino Uno R4 WiFi", "5 V", "100 mA", "200 mA", "0.5 W", "1 W", "USB / ESC BEC"],
        ["Steering Servo", "5 V", "50 mA", "500 mA", "0.25 W", "2.5 W", "ESC BEC"],
        ["MSC-25RC ESC", "7.2 V", "0.14 to 0.28 A", "3.5 A+", "1 to 2 W", "25 W+", "7.2 V LiPo"],
        ["Drive Motor", "7.2 V", "0.4 to 0.7 A", "4 A+", "3 to 5 W", "30 W+", "7.2 V LiPo"],
        ["Car Group Total", "", "0.7 to 1.1 A", "8.2 A+", "5 to 8 W", "59 W", ""],
        ["System Total", "", "2.7 to 3.7 A", "12.2 A+", "18 to 25 W", "85 W", ""],
    ],
    [1.5, 0.8, 1.1, 1.0, 1.0, 0.9, 1.0]
)
add_caption(doc, "Table 3.3: System power requirements by group.")

add_heading(doc, "3.5 Phase 2 Software Stack", 2)
add_para(doc, "The software stack is distributed across the three compute platforms as follows:")
add_bullet(doc, "Jetson Orin Nano: Ubuntu 22.04, JetPack 6, Python 3.10, PyQt5, PyTorch with CUDA, OpenCV 4.8.0, Ultralytics YOLOv8")
add_bullet(doc, "Raspberry Pi 5: Python 3.11, Infineon ifxAvian SDK providing the BGT60TR13C radar interface utilities")
add_bullet(doc, "Arduino Uno R4 WiFi: C++ firmware using the Arduino Servo library for 50 Hz PWM generation")
add_para(doc, "All inter-device communication uses standard UDP sockets over the shared WiFi hotspot with no additional middleware [22].")

add_heading(doc, "3.6 Phase 2 Distributed Architecture and Threading", 2)
add_para(doc, "The system distributes work across three boards over a shared WiFi hotspot. The Pi runs the radar pipeline and streams confirmed track data to the Jetson on UDP port 9576. The Jetson runs lidar parsing, YOLOv8n inference, sensor fusion, motor decisions and data logging. The Arduino translates UDP commands into ESC and servo PWM. The Arduino Uno R4 WiFi is configured with a static IP address of 172.20.10.3 set using WiFi.config before WiFi.begin, ensuring a consistent address on every power cycle without router configuration [22]. The Jetson also listens on a separate UDP port for pause and resume commands from a laptop, allowing immediate halt during testing.")
add_para(doc, "The fusion application runs four concurrent threads: the lidar thread reads the LD06 serial port and updates shared state under a threading.Lock; the radar thread listens on UDP port 9576 and updates shared radar state as JSON packets arrive; the camera thread captures 416 by 416 frames through GStreamer nvarguscamerasrc with nvvidconv hardware resize and runs YOLOv8n on CUDA every second frame to manage thermal load; the GUI thread polls shared state every 100 ms via QTimer and re-renders the dashboard without waiting on sensor data, sustaining approximately 9.5 FPS [22]. Real test logs confirm the fusion loop running at approximately 15 Hz with consistent 60 to 65 ms timestamps and no UDP lag.")
add_para(doc, "Figure 3.4 shows the fusion GUI dashboard during operation.")
add_image(doc, "fig3_4_fusion_gui", "Figure 3.4: Fusion GUI showing lidar scatter panel, radar range-velocity panel, TTC chart and status bar.")
add_para(doc, "As shown in Figure 3.4, the left panel renders the lidar bird's-eye scatter plot coloured by zone, the centre panel shows confirmed radar tracks on a range-velocity plot coloured by alert level, the right panel contains the TTC bar chart and track table, and the bottom status bar provides a single-line snapshot of FPS, sensor counts, fused verdict and health state.")

add_heading(doc, "3.7 Phase 2 Sensor Fusion and Health Monitoring", 2)
add_para(doc, "The fusion policy is worst-case-across-sensors: Imminent maps to priority 3, Caution to 2 and Safe to 1. The highest priority sensor governs the output and no sensor can reduce the result [11], [13]. Lidar handles stationary obstacle detection. Radar handles moving targets through TTC. The camera cross-validates lidar alerts: if YOLO finds nothing in the central 40 percent of the frame the lidar alert is suppressed as noise, and if both agree the detection is confirmed [12]. The camera cannot trigger alerts independently. The health monitor floors the output at Caution if both lidar and camera are simultaneously impaired ensuring the system never returns Safe when two primary sensors are unavailable [13].")
add_para(doc, "The health layer distinguishes a disconnected sensor from a degraded one. A disconnected sensor stops sending data and its timestamp goes stale. A degraded sensor is still transmitting but cannot detect anything useful: a lidar with a blocked lens continues sending packets with zero valid forward-zone points, or a camera in a scene too dark for detection [3], [4]. Lidar degradation is flagged after 15 consecutive frames of zero forward-zone points and triggers radar static fallback coverage. Camera degradation triggers when mean brightness falls below 40 and current detections are cleared immediately.")

add_heading(doc, "3.8 Phase 2 LD06 Lidar Pipeline", 2)
add_para(doc, "The LD06 transmits 47-byte packets at 230400 baud with 12 distance measurements per packet. Bytes 4 to 5 encode the start angle and bytes 42 to 43 encode the end angle, both divided by 100 to convert to degrees, with each measurement occupying 3 bytes for distance low byte, distance high byte and confidence. When first connected every obstacle appeared behind the origin because the sensor native zero-degree reference points opposite to the vehicle forward direction. An angle offset correction applied as (angle + ANGLE_OFFSET_DEG) % 360.0 and a Y-axis sign inversion of minus 1 resolved this [1], [2]. Three filtering layers were then applied in sequence:")
add_bullet(doc, "FOV restriction to plus or minus 20 degrees forward, a 20 cm minimum distance floor and a secondary geometric check for near-origin stray points [2]")
add_bullet(doc, "Minimum point count per zone scaled with detection distance so closer objects need more points to register than distant ones [3]")
add_bullet(doc, "Persistence filter requiring 3 consecutive scans to raise an alert and 4 consecutive clear scans to lower it, with asymmetric hysteresis preventing state flickering [4]")
add_para(doc, "These three layers removed virtually all false positives observed during bench testing. Two additional threshold changes were applied based on evidence from field test logs. The centre zone width was reduced from 19 cm to 12 cm total because a bottle is approximately 8 cm wide and the wider zone caused thin objects placed slightly off-centre to register in both the centre and left zones simultaneously. The stop threshold was raised from 0.35 m to 0.6 m to match real sensor readings: the lidar was consistently reading 0.44 to 0.48 m when physically blocked at close range and with the original 0.35 m threshold the reverse manoeuvre was never triggering [22]. Table 3.4 defines the three detection zones.")
add_table(doc,
    ["Zone", "Angular Range", "Navigation Response"],
    [
        ["Centre (red)", "340 to 360 and 0 to 20 degrees", "Reduce speed or stop"],
        ["Left (blue)", "20 to 90 degrees", "Steer right to avoid"],
        ["Right (green)", "270 to 340 degrees", "Steer left to avoid"],
    ],
    [2.0, 3.0, 2.0]
)
add_caption(doc, "Table 3.4: Lidar zone definitions and navigation responses.")
add_para(doc, "Figure 3.5 shows the corrected live scatter plot with points coloured by zone and threshold lines overlaid.")
add_image(doc, "fig3_5_lidar_scatter", "Figure 3.5: LD06 live scatter plot with three-zone colouring and threshold lines.")
add_para(doc, "As shown in Figure 3.5, the centre zone in red, left in blue and right in green are clearly distinguishable in the bird's-eye view with the two horizontal threshold lines marking the stop and slow distances providing a real-time reference for when detections are entering the danger range.")
add_github(doc, "Phase2 Hardare Python Scripts/initial_Test_Debug_Scripts/lidar_GUI.py")

add_heading(doc, "3.9 Phase 2 BGT60TR13C Radar Pipeline", 2)
add_para(doc, "The BGT60TR13C runs on the Pi 5 with one transmit and three receive antennas producing a raw data array of shape (3, 64, 128) per frame. Five processing layers are applied in sequence [5], [7], [8]:")
add_bullet(doc, "Three-stage FFT: Range FFT along the sample axis gives distance, Doppler FFT along the chirp axis gives velocity and Angle FFT across the three receivers gives bearing")
add_bullet(doc, "MTI background subtraction at memory coefficient 0.92 removes static clutter from walls, floors and the chassis, requiring approximately 15 frames to settle [5]")
add_bullet(doc, "CA-CFAR detection at 16 dB identifies candidates above the adaptive local noise floor [6]")
add_bullet(doc, "Nearest-neighbour tracking confirms tracks after 3 consecutive matched frames and deletes after 3 consecutive misses")
add_bullet(doc, "Approach confirmation gate requiring 3 consecutive frames of consistently negative radial velocity before a track triggers a warning [7]")
add_para(doc, "TTC triggers Imminent below 1.5 seconds or within 0.5 m and Caution below 3.0 seconds or within 1.2 m, with range fallbacks for near-zero velocity [14], [15]. Stationary obstacles are the lidar's responsibility entirely because MTI suppresses static returns by design.")
add_para(doc, "Figure 3.6 shows the radar GUI with a confirmed approaching track displayed on the range-velocity plot.")
add_image(doc, "fig3_6_radar_gui", "Figure 3.6: Radar GUI showing confirmed approaching track with range-velocity position and TTC.")
add_para(doc, "As shown in Figure 3.6, confirmed tracks appear as dots at their range and velocity coordinates coloured by alert level with TTC values annotated alongside each track, allowing the operator to see at a glance which targets are approaching and how urgently.")
add_github(doc, "Phase2 Hardare Python Scripts/radar_send.py")

add_heading(doc, "3.10 Phase 2 Camera Pipeline, Motor Control and Avoidance State Machine", 2)
add_para(doc, "The Arducam IMX477 connects via CSI-2 through a GStreamer nvarguscamerasrc pipeline hardware-resizing to 416 by 416 via nvvidconv. YOLOv8n runs on CUDA every second frame. Only detections whose bounding box centre x falls within the middle 40 percent of the frame are considered. Distance estimation uses projective geometry: estimated distance equals real-world object height multiplied by focal length divided by bounding box pixel height, producing an Imminent, Caution or Safe band without a depth sensor [10]. Camera verdicts are cleared when mean brightness falls below 40. A significant finding from real-world testing was that the IMX477 ribbon cable hanging loose on the front-left side of the chassis entered the lidar scan plane, producing a permanent false left-zone reading of 0.33 to 0.40 m in open space and triggering continuous Imminent state. The fix was to route the cable behind the lidar and raise the sensor on standoffs [22].")
add_para(doc, "Figure 3.7 shows the Arducam IMX477 connected to the Jetson via the CSI-2 port.")
add_image(doc, "fig3_7_camera", "Figure 3.7: Arducam IMX477 connected to Jetson via CSI-2.")
add_para(doc, "As shown in Figure 3.7, the camera connects via ribbon cable directly to the CSI-2 port on the Jetson Orin Nano. The routing of this cable away from the lidar scan plane was identified as a critical hardware fix following real-world testing.")
add_github(doc, "Phase2 Hardare Python Scripts/initial_Test_Debug_Scripts/yolo_test.py")
add_para(doc, "Motor commands are encoded as JSON and sent from the Jetson to the Arduino on UDP port 5005. Safe commands 1700 microseconds forward, Caution commands 1570 microseconds slow forward and Imminent commands 1500 microseconds neutral stop. Steering derives from lidar zone output: a left zone obstacle triggers right steering, a right zone obstacle triggers left steering and both zones clear returns the servo to 1500 microseconds centre.")
add_para(doc, "The motor controller implements a six-state avoidance machine running at 20 Hz. The key design principle is sensor-driven state transitions rather than fixed timers: states exit as soon as sensor conditions are met rather than running for a predetermined duration. Reversing is the only state retaining a fixed duration of 2.5 seconds to guarantee physical clearance before a turn is attempted. Normal operates across three distance zones: above 1.5 m the car drives at full forward speed, between 1.0 m and 1.5 m it slows and steers gently away from the obstacle, and between 0.6 m and 1.0 m it steers around the obstacle at slow speed without stopping. Only when an obstacle is under 0.6 m and both sides are simultaneously blocked does the car stop and reverse. Stopping holds full brake and selects an escape direction. Reversing drives straight reverse for 2.5 seconds. Turning applies slow forward toward the escape direction and exits the moment the front centre distance reads above 1.0 m. Escaping drives slow forward on the new heading, replanning if a new obstacle appears after a 0.5 second grace period, and straightening temporarily if a side wall falls within 0.25 m. Straightening recentres steering and resets retry counters. A stuck detection timer handles scenarios where the car has been continuously steering around an obstacle for more than 4 seconds without clearing the front [22].")
add_para(doc, "Figure 3.8 shows the updated sensor-driven avoidance state machine diagram.")
add_image(doc, "fig3_8_statemachine", "Figure 3.8: Sensor-driven avoidance state machine with six states, sensor-driven exits and retry logic.")
add_para(doc, "As shown in Figure 3.8, the state machine shows all six states with their timing constraints, sensor-driven exit conditions, and the retry and direction flip paths that activate when the chosen escape direction is repeatedly blocked.")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/motor_controller.py")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/car_control.py")

add_heading(doc, "3.11 Phase 2 ESC Reverse Arming and Autostart", 2)
add_para(doc, "The MSC-25RC ESC requires a specific three-step arming sequence before reverse engages: a brake pulse at 1300 microseconds for 200 ms, a neutral pulse at 1500 microseconds for 300 ms, then the reverse pulse at 1300 microseconds, taking approximately 500 ms total. The motor controller was sending the reverse command every 50 ms at 20 Hz, restarting the arming sequence each time so the car never physically reversed. A reverseArmed flag on the Arduino fixed this: once the sequence completes all subsequent reverse commands are ignored and the ESC holds reverse PWM continuously until a different command arrives [22].")
add_para(doc, "Figure 3.9 shows the ESC reverse arming sequence and the reverseArmed flag logic.")
add_image(doc, "fig3_9_esc_arming", "Figure 3.9: ESC reverse arming sequence showing IDLE to BRAKE to NEUTRAL WAIT to REVERSING with reverseArmed flag.")
add_para(doc, "As shown in Figure 3.9, without the reverseArmed flag the arming sequence restarts on every incoming command and the ESC never reaches the Reversing state. With the flag in place the sequence completes once and reverse is maintained until a different motor command arrives.")
add_para(doc, "Both boards start automatically on power-on: the radar streamer on the Pi runs as a systemd service and the Jetson launches the fusion GUI via a desktop autostart file, requiring no laptop for field operation.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "Chapter 4: Performance Evaluation Setup and Scenarios", 1)

add_heading(doc, "4.1 Evaluation Methodology", 2)
add_para(doc, "The test structure follows a progressive validation methodology. Phase 1 establishes a simulation baseline through three scenarios of increasing complexity: a clear conditions single-obstacle test to validate the core pipeline under ideal conditions, a rain variant holding all other variables constant to isolate sensor redundancy by stressing the camera, and a multi-obstacle sequential test to verify state machine integrity over repeated encounters. This mirrors controlled ablation study design where one variable changes at a time to attribute results to specific system components [11], [13]. Phase 2 bench testing preceded on-car testing as a safety precondition to localise faults to the integration layer when they appear. The two on-car sessions are structured as an extended multi-obstacle stress test and a single-obstacle controlled reproduction of the Phase 1 baseline, allowing direct comparison between simulation and hardware performance.")
add_para(doc, "Metrics assessed across both phases are detection rate as true positives over total obstacles encountered, false positive rate as spurious alerts in clear conditions, TTC accuracy as estimated versus measured approach time, reaction latency from fusion alert to first motor command, stop distance at halt, steering accuracy, drive time efficiency as proportion of session time in normal forward driving and avoidance success rate as the proportion of encounters completing the full manoeuvre sequence. Phase 1 metrics come from the CSV log files and Phase 2 metrics from the on-board data logger.")

section_banner(doc, "Phase 1 Evaluation: Carla Simulation Scenarios", "1A6E3C")

add_heading(doc, "4.2 Phase 1 Test Environment", 2)
add_para(doc, "All Phase 1 runs execute in Carla on Town04 in synchronous mode at 20 Hz, reset to a clean state between runs. Ego vehicle spawn, obstacle placement and threshold constants are identical across all runs so the only variable changing between scenarios is weather conditions and obstacle count [20], [21]. In Phase 1 each scenario script constitutes both the test configuration and the test execution. Running the script under the defined conditions produces the logged CSV outputs from which all Phase 1 metrics are derived. This differs from Phase 2 where the test environment is physical and results reflect real-world sensor behaviour rather than simulated sensor models.")

add_heading(doc, "Scenario 1: Single Obstacle, Clear Conditions", 2)
add_heading(doc, "4.3 Scenario 1 Setup", 3)
add_para(doc, "A single stationary obstacle vehicle is placed 80 metres ahead in the ego lane under default clear weather with all three simulated sensors operating normally. This is the clean pipeline baseline against which all other scenario results are compared. The ego vehicle approaches from the spawn point at normal throttle following the waypoint lane track. As the fused distance drops below 15 m the system enters Warning and blends avoidance steering with lane following. When fused distance drops below 8 m or TTC drops below 2 seconds the system transitions to Critical and applies full emergency braking. After the vehicle stops, Lane Change is initiated with left steer at minus 0.25 and throttle 0.35. Straightening then blends back to lane steering at throttle 0.45 and Resumed restores normal driving at throttle 0.5 [20], [22].")
add_para(doc, "Figure 4.1 shows the Town04 initial setup with the ego vehicle and stationary obstacle aligned in the driving lane.")
add_image(doc, "fig4_1_town04", "Figure 4.1: Town04 initial setup with ego vehicle and obstacle aligned in lane with sensor overlays.")
add_para(doc, "As shown in Figure 4.1, the simulation provides a clear straight approach path with the obstacle placed directly ahead in the same lane, giving the ego vehicle sufficient forward distance to detect and respond before reaching the critical distance threshold.")
add_para(doc, "Figure 4.2 shows the emergency brake trigger moment under clear conditions.")
add_image(doc, "fig4_2_brake_clear", "Figure 4.2: Emergency brake trigger moment under clear conditions.")
add_para(doc, "As shown in Figure 4.2, the vehicle has transitioned to Critical state with the fused distance having dropped below 8 m. Throttle is set to zero and full braking is applied, bringing the vehicle to a complete stop before contact with the obstacle.")
add_para(doc, "Figure 4.3 shows the lane change manoeuvre beginning after the vehicle has stopped.")
add_image(doc, "fig4_3_lane_change", "Figure 4.3: Lane change manoeuvre, Scenario 1.")
add_para(doc, "As shown in Figure 4.3, the vehicle applies gentle left steering at minus 0.25 with throttle 0.35 to begin moving into the adjacent lane, clearing the obstacle without reversing.")
add_para(doc, "Figure 4.4 shows the vehicle straightening after completing the lane change.")
add_image(doc, "fig4_4_straightening", "Figure 4.4: Vehicle straightening after lane change.")
add_para(doc, "As shown in Figure 4.4, the steering is blending back toward the lane-following signal at throttle 0.45 as the vehicle realigns with the new lane ahead of the obstacle.")
add_para(doc, "Figure 4.5 shows the vehicle resuming normal lane driving.")
add_image(doc, "fig4_5_resume", "Figure 4.5: Vehicle resuming normal lane driving.")
add_para(doc, "As shown in Figure 4.5, the vehicle has returned to the Resumed state with full throttle at 0.5 and complete waypoint-guided steering restored, driving normally past the cleared obstacle.")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front.py")

add_heading(doc, "4.4 Scenario 1 Results and Analysis", 3)
add_para(doc, "Radar readings fluctuated around 12 to 14 m during the approach, reflecting characteristics of the Carla radar simulation model. Lidar was considerably more stable and settled at approximately 4.3 m at closest range. The fused distance correctly overrode the noisier radar readings through the 12 m association gate throughout the approach [11]. Both the distance and TTC thresholds converged at the same braking event, confirming the dual-trigger logic is internally consistent and neither condition is redundant. Safety state distribution across the 70-second run was approximately 89 percent Safe, 2 to 3 percent Warning and 8 percent Critical. The radar to lidar mean distance difference of approximately 8 to 9 m was absorbed correctly by the association gate. TTC values remained above 2 seconds throughout the run with the sub-threshold cluster appearing exclusively at the braking event. Lane change, straightening and resume completed without collision. Scenario 1 is a pass.")
add_para(doc, "Figure 4.6 shows the performance dashboard generated from the Scenario 1 log files.")
add_image(doc, "fig4_6_dashboard", "Figure 4.6: Scenario 1 performance dashboard showing sensor distances over time, TTC profile and safety state distribution.")
add_para(doc, "As shown in Figure 4.6, the sensor distances over time graph confirms radar fluctuating above lidar during the approach with the fused distance correctly tracking the lower validated reading. The TTC graph shows a sharp drop through the 2 second threshold at the braking event. The safety state pie chart confirms the dominant Safe proportion at approximately 89 percent with the small Critical region corresponding to the braking and post-stop period. The radar to lidar distance agreement histogram confirms the mean difference of approximately 8 to 9 m which the association gate absorbed correctly throughout the run.")

add_heading(doc, "Scenario 2: Single Obstacle, Extreme Rain", 2)
add_heading(doc, "4.5 Scenario 2 Setup", 3)
add_para(doc, "Scenario 2 repeats the identical approach, braking, lane change and resume sequence as Scenario 1 with the only variable being extreme adverse weather applied through Carla's WeatherParameters API. Precipitation is set to 100 percent, road wetness to 100 percent, fog density to 40 percent with onset at 30 m and sun altitude to minus 30 degrees below the horizon. Headlights are enabled automatically. The purpose of Scenario 2 is to isolate sensor redundancy: if the braking and recovery sequence matches Scenario 1 despite heavy camera degradation it proves that lidar and radar alone are driving the safety decision independently of camera state [20], [22].")
add_table(doc,
    ["Parameter", "Scenario 1 Clear", "Scenario 2 Extreme Rain"],
    [
        ["Precipitation", "None", "100 percent"],
        ["Precipitation deposits", "None", "100 percent, wet roads"],
        ["Sun altitude angle", "Default", "Minus 30 degrees, below horizon"],
        ["Fog density", "None", "40 percent"],
        ["Fog start distance", "Not applicable", "30 m"],
        ["Road wetness", "Dry", "100 percent"],
        ["Headlights", "Off", "On"],
    ],
    [2.5, 2.0, 2.5]
)
add_caption(doc, "Table 4.1: Weather parameter comparison between Scenario 1 clear and Scenario 2 extreme rain.")
add_para(doc, "Figure 4.7 shows the extreme rain simulation environment.")
add_image(doc, "fig4_7_rain_setup", "Figure 4.7: Extreme rain setup showing fog, wet roads and reduced visibility.")
add_para(doc, "As shown in Figure 4.7, the combination of fog starting at 30 m and the sun altitude set to minus 30 degrees below the horizon creates substantially reduced forward camera visibility. The wet road surface and precipitation overlay are clearly visible across the scene.")
add_para(doc, "Figure 4.8 shows the braking event under extreme rain conditions.")
add_image(doc, "fig4_8_brake_rain", "Figure 4.8: Braking under extreme rain conditions.")
add_para(doc, "As shown in Figure 4.8, despite the heavily degraded camera stream the system correctly transitions to Critical state and applies full emergency braking at the same fused distance threshold as Scenario 1, driven by lidar and radar independently of camera state.")
add_para(doc, "Figure 4.9 shows the lane change and resume under extreme rain.")
add_image(doc, "fig4_9_lanechange_rain", "Figure 4.9: Lane change and resume under extreme rain.")
add_para(doc, "As shown in Figure 4.9, the lane change, straightening and resume sequence completes correctly under rain conditions with no difference in behaviour compared to the clear conditions baseline, confirming the planned post-stop recovery is unaffected by weather.")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front_rain.py")

add_heading(doc, "4.6 Scenario 2 Results and Analysis", 3)
add_para(doc, "The camera stream was heavily degraded throughout the run. Fog onset at 30 m cut forward visible range substantially and the precipitation overlay reduced scene contrast across the frame. Lidar distance extraction remained stable throughout because the ray-cast lidar model in Carla does not simulate precipitation attenuation, consistent with real lidar behaviour in light to moderate rain. Radar readings were equally unaffected by visibility conditions [8]. The fusion pipeline produced the same Warning and Critical transitions at the same threshold values as Scenario 1 with no meaningful delay. This demonstrates directly that the multi-sensor architecture maintains safety-critical detection under conditions where a camera-only system would struggle substantially at range. The lane change and recovery sequence completed correctly under rain. Scenario 2 is a pass.")

add_heading(doc, "Scenario 3: Sequential Multi-Obstacle", 2)
add_heading(doc, "4.7 Scenario 3 Setup", 3)
add_para(doc, "Scenario 3 tests whether the safety state machine resets correctly between successive avoidance cycles without state contamination from previous encounters. Three obstacle vehicles are configured with the first placed at 150 m ahead in the driving lane. Each subsequent obstacle is spawned dynamically 150 m ahead in the current driving lane as the preceding one is passed. Obstacles are assigned distinct colours: red, green and blue respectively for visual identification in the output video. An obstacle is classified as passed when the forward dot product of the ego-to-obstacle vector becomes negative or zero, at which point the braked and lane_change_complete flags reset and the next obstacle is treated as a completely fresh event [20], [22].")
add_github(doc, "Phase1 Carla Project Python Scripts/Initial_Tests_Scripts/dt_phase1_car_front_multi_obs_.py")

add_heading(doc, "4.8 Scenario 3 Results and Analysis", 3)
add_para(doc, "All three obstacles spawned correctly in the driving lane at 150 m ahead. Each triggered a fully independent detection, Warning transition, Critical transition, emergency brake and lane change sequence with no evidence of state contamination from the previous cycle. The dot product classification correctly identified each obstacle as passed once cleared and the flag reset allowed the state machine to respond to each subsequent obstacle as a fresh event. Between cycles the vehicle returned cleanly to lane-following at throttle 0.5 with no residual steering anomaly. The output video Multi Object Avoidance.mp4 confirms all three avoidance cycles visually. Scenario 3 is a pass.")

add_heading(doc, "4.9 Phase 1 YOLO Performance Evaluation", 2)
add_heading(doc, "4.9.1 YOLOv8n Baseline Evaluation", 3)
add_para(doc, "YOLOv8n was evaluated offline over the recorded camera frames from all Phase 1 scenarios using Live_Visualisation_with_YOLO.py. Mean inference time was approximately 9.2 ms per frame corresponding to approximately 114 FPS after GPU warm-up, with an initial spike of approximately 900 ms at the first frame for model loading and GPU memory allocation. Average in-lane detection count was approximately 0.12 per frame with car as the dominant class, consistent with the Town04 obstacle type. These results confirm YOLO is not a latency bottleneck in the Phase 1 pipeline [9].")

add_heading(doc, "4.9.2 YOLOv5n versus YOLOv8n Comparative Evaluation", 3)
add_para(doc, "As an extension to the baseline YOLO evaluation, YOLOv5n and YOLOv8n were compared across both clear and rain scenario conditions to determine which model provides superior real-time performance. Both models used the identical lane-relevance filtering pipeline across 2,331 frames for the clear conditions scenario and 1,185 frames for the rain scenario [9], [22].")
add_para(doc, "Under clear conditions YOLOv8n achieved a mean inference time of 19.8 ms against 20.5 ms for YOLOv5n, corresponding to mean FPS of 50.6 versus 49.4. YOLOv5n produced a catastrophic spike of 673.90 ms during the clear conditions run, representing a 33 times slowdown from its typical performance. At 10 m/s vehicle speed this creates an estimated 6.7 m blind spot. YOLOv8n's worst-case inference time was 67.88 ms, 10 times more stable, creating only a 0.68 m blind spot under equivalent conditions. Standard deviation of inference time was 13.56 ms for YOLOv5n against 1.18 ms for YOLOv8n. Car detections were 133 for YOLOv8n versus 129 for YOLOv5n [9].")
add_para(doc, "Under extreme rain conditions YOLOv5n detected 28 cars versus 21 for YOLOv8n. YOLOv8n produced 7 unique detected classes versus 5 for YOLOv5n, with the higher class count reflecting more misclassifications under degraded frames. Table 4.2 summarises the full comparison.")
add_table(doc,
    ["Metric", "YOLOv5n Clear", "YOLOv8n Clear", "YOLOv5n Rain", "YOLOv8n Rain"],
    [
        ["Car detections", "129", "133", "28", "21"],
        ["Mean inference time", "20.5 ms", "19.8 ms", "19.9 ms", "19.7 ms"],
        ["Mean FPS", "49.4", "50.6", "50.3", "51.0"],
        ["Min FPS", "1.5", "10.9", "14.5", "17.7"],
        ["Max inference time", "673.90 ms", "67.88 ms", "68.93 ms", "56.49 ms"],
        ["Std deviation", "13.56 ms", "1.18 ms", "1.55 ms", "not stated"],
        ["Unique classes", "1", "1", "5", "7"],
    ],
    [2.0, 1.4, 1.4, 1.4, 1.4]
)
add_caption(doc, "Table 4.2: YOLOv5n versus YOLOv8n performance comparison under clear and rain conditions.")
add_para(doc, "YOLOv8n was selected for the final pipeline on the basis of its significantly superior inference stability under clear conditions. The catastrophic single-frame spike produced by YOLOv5n is unacceptable for a real-time collision avoidance system [9]. Figure 4.10 shows the YOLO performance dashboards for both models.")
add_para(doc, "As shown in Figure 4.10, the inference time graphs confirm the catastrophic 673.90 ms spike in YOLOv5n under clear conditions, which is clearly absent in the YOLOv8n equivalent graph. The FPS graphs show YOLOv5n dropping to 1.5 FPS at the spike point while YOLOv8n maintains a minimum of 10.9 FPS.")
add_github(doc, "Phase1 Carla Project Python Scripts/Visualisation Scripts/Live_Visualisation_with_YOLO.py")
add_github(doc, "Phase1 Carla Project Python Scripts/Visualisation Scripts/analyse_logs.py")

section_banner(doc, "Phase 2 Evaluation: On-Car Testing", "2E75B6")

add_heading(doc, "4.10 Phase 2 Test Environment", 2)
add_para(doc, "All Phase 2 testing was conducted with the system running fully autonomously on the Maverick RC car across both indoor and outdoor environments. The on-board data logger writes a row only when an obstacle is detected or the motor state changes, capturing lidar zone distances for all three zones, radar range and velocity, YOLO detections, fusion level and active motor state per row [22]. Two test sessions were conducted on 1 April 2026.")

add_heading(doc, "Test 1: Multi-Obstacle Extended Run", 2)
add_heading(doc, "4.11 Test 1 Setup", 3)
add_para(doc, "Test 1 placed multiple obstacles at various distances and angles across the car's path to force repeated avoidance manoeuvres over an extended autonomous run. The goal was to stress-test the retry and direction flip logic across multiple consecutive encounters and evaluate whether the system degrades over time under sustained operation. The session ran for approximately 1000 seconds and produced 15,243 logged events. Obstacles were placed at varying angles and distances to create a range of approach geometries covering centre-blocked, side-blocked and combined scenarios.")

add_heading(doc, "4.12 Test 1 Results and Analysis", 3)
add_para(doc, "Figure 4.11 shows the lidar zone distances over time for Test 1 with motor state background shading.")
add_image(doc, "fig4_11_lidar_time", "Figure 4.11: Lidar zone distances over time for Test 1 with motor state shading.")
add_para(doc, "As shown in Figure 4.11, the forward centre and right zone distances drop repeatedly to near zero at each obstacle encounter before recovering as the car completes its avoidance manoeuvre. The left zone shows a consistent low reading throughout the session at approximately 0.33 to 0.45 m, which is the ribbon cable false reading rather than a real left-zone obstacle. The coloured background shading confirms the motor state active at each point, showing Braking, Reversing and Escaping states aligned with the distance drops.")
add_para(doc, "Figure 4.12 shows the motor state distribution for Test 1.")
add_image(doc, "fig4_12_motor_states", "Figure 4.12: Motor state distribution for Test 1.")
add_para(doc, "As shown in Figure 4.12, Normal state dominates at 93.4 percent of all logged events with the avoidance states of Braking, Reversing, Turning and Escaping collectively accounting for the remaining 6.6 percent, confirming the system operates in forward driving mode for the large majority of the session and that avoidance manoeuvres are handled efficiently without prolonged recovery periods.")
add_para(doc, "Figure 4.13 shows the fusion level distribution for Test 1 as a pie chart with a timeline below.")
add_image(doc, "fig4_13_fusion_levels", "Figure 4.13: Fusion level pie chart and timeline for Test 1.")
add_para(doc, "As shown in Figure 4.13, the pie chart shows Imminent at 61.2 percent, Caution at 34.8 percent and Safe at only 3.9 percent. The timeline makes the cause clear: the fusion level rarely drops out of Imminent or Caution throughout the session, which reflects the ribbon cable false reading keeping the left zone permanently triggered rather than genuine continuous obstacle presence. In a correctly mounted configuration the Safe proportion would dominate during open driving.")
add_para(doc, "Figure 4.14 shows the YOLO detection classes for Test 1.")
add_image(doc, "fig4_14_yolo_det", "Figure 4.14: YOLO detection classes for Test 1.")
add_para(doc, "As shown in Figure 4.14, the detections are spread across everyday object classes including bottle, vase, book and laptop, reflecting the nearest COCO class match for the physical test obstacles rather than formal traffic hazard categories. The low overall YOLO coverage of 14 percent in Test 1 is consistent with the car spending much of the extended session driving between obstacle encounters where nothing is visible in the forward camera zone. Radar shows zero detections as expected because all test obstacles were stationary and correctly suppressed by the MTI filter [5].")
add_para(doc, "Figure 4.15 shows the lidar distance histograms per zone for Test 1.")
add_image(doc, "fig4_15_lidar_hist", "Figure 4.15: Lidar distance histograms per zone for Test 1.")
add_para(doc, "As shown in Figure 4.15, the left zone histogram shows a sharp spike concentrated between 0.33 m and 0.45 m across the entire session, confirming the ribbon cable as a persistent false reading in that zone. The centre and right zone histograms show a natural spread with peaks at various distances reflecting real obstacle encounters. The stop threshold line at 0.6 m and steer threshold at 0.8 m are marked on each histogram.")
add_para(doc, "Figure 4.16 shows the avoidance event summary for Test 1.")
add_image(doc, "fig4_16_avoidance", "Figure 4.16: Avoidance event summary for Test 1.")
add_para(doc, "As shown in Figure 4.16, the summary table confirms 15 braking events across the session with 13 progressing all the way through to Escaping, giving an 86.7 percent avoidance success rate. The 2 incomplete events are identified as either mid-sequence direction flips when the chosen escape path was also blocked, or natural clears before reversing was required, rather than system failures.")

add_heading(doc, "Test 2: Single Obstacle Controlled Run", 2)
add_heading(doc, "4.13 Test 2 Setup", 3)
add_para(doc, "Test 2 placed a single everyday object directly ahead in the forward centre zone to validate a clean detection and escape cycle in an uncluttered environment. The goal was a controlled reproduction of the Phase 1 single-obstacle baseline on real hardware, providing a direct comparison between simulation and physical system performance under equivalent conditions. The session ran for approximately 30 seconds and produced 985 logged events.")

add_heading(doc, "4.14 Test 2 Results and Analysis", 3)
add_para(doc, "The right zone shows a sustained low reading consistent with the single obstacle placed to the right of the car's path. The avoidance event is clearly visible as a brief Braking and Reversing period before the car transitions back to Normal, with the distance recovering after the car escapes to a new heading. Test 2 produced a single braking event which progressed successfully all the way through to Escaping and Straightening, giving a 100 percent avoidance success rate for the single-obstacle controlled test. YOLO coverage was 53 percent in Test 2, higher than Test 1, consistent with more of the shorter session spent in close proximity to the single obstacle. Radar shows zero detections as expected for a stationary obstacle [5].")

add_heading(doc, "4.15 Combined System Performance Metrics", 2)
add_para(doc, "Figure 4.17 shows the overall system performance indicators as percentages for direct comparison between both test sessions.")
add_image(doc, "fig4_17_perf", "Figure 4.17: System performance metrics comparing Test 1 and Test 2.")
add_para(doc, "As shown in Figure 4.17, the side-by-side comparison highlights the consistency between sessions in drive time efficiency at 93 percent for both, while showing the impact of the ribbon cable on lidar coverage at 87.2 percent in Test 1 versus 100 percent in Test 2. YOLO coverage was 14 percent in Test 1 and 53 percent in Test 2. Radar coverage was 0 percent in both sessions as expected for stationary obstacles. The obstacle response rate of 93.3 percent in Test 1 confirms the car committed to the complete avoidance manoeuvre in 14 of the 15 braking events. The avoidance success rate difference of 86.7 percent in Test 1 against 100 percent in Test 2 reflects the additional complexity of the multi-obstacle extended environment rather than a fundamental system failure, as the single-obstacle clean result confirms the core pipeline is functioning correctly. The most significant outstanding fix identified from both test sessions is the ribbon cable routing: once resolved the fusion level distribution will reflect actual obstacle encounters rather than a permanent sensor impairment [22].")
add_github(doc, "Phase2 Hardare Python Scripts/jetson_fusion.py")
add_github(doc, "Phase2 Hardare Python Scripts/Motor scripts/motor_controller.py")

doc.add_page_break()

# ── REFERENCES ────────────────────────────────────────────────────────────────
add_heading(doc, "References", 1)
refs = [
    '[1] T. Raj, F. H. Hashim, A. B. Huddin, M. F. Ibrahim and A. Hussain, "A Survey on LiDAR Scanning Mechanisms," Electronics, vol. 9, no. 5, p. 741, 2020.',
    '[2] R. Roriz, J. Cabral and T. Gomes, "Automotive LiDAR Technology: A Survey," IEEE Trans. Intell. Transp. Syst., vol. 23, no. 7, pp. 6282-6297, 2022.',
    '[3] S. Shi et al., "3D Point Cloud Processing and Learning for Autonomous Driving," arXiv preprint arXiv:2003.00601, 2020.',
    '[4] Y. Li, C. Zhao, S. Liu and J. Wang, "Algorithm for Point Cloud Dust Filtering of LiDAR for Autonomous Vehicles in Mining Area," Sustainability, vol. 16, no. 7, p. 2827, 2024.',
    '[5] M. A. Richards, Fundamentals of Radar Signal Processing, 2nd ed. New York, NY: McGraw-Hill, 2014.',
    '[6] H. Rohling, "Radar CFAR Thresholding in Clutter and Multiple Target Situations," IEEE Trans. Aerosp. Electron. Syst., vol. 19, no. 4, pp. 608-621, 1983.',
    '[7] X. Huang et al., "Exploring Radar Data Representations in Autonomous Driving: A Comprehensive Review," arXiv preprint arXiv:2312.04861, 2024.',
    '[8] I. Bilik, O. Longman, S. Villeval and J. Tabrikian, "The Rise of Radar for Autonomous Vehicles," IEEE Signal Process. Mag., vol. 36, no. 5, pp. 20-31, 2019.',
    '[9] J. Redmon, S. Divvala, R. Girshick and A. Farhadi, "You Only Look Once: Unified, Real-Time Object Detection," in Proc. IEEE CVPR, Las Vegas, NV, USA, 2016, pp. 779-788.',
    '[10] M. Vajgl and P. Hurtik, "Dist-YOLO: Fast Object Detection with Distance Estimation," Applied Sciences, vol. 12, no. 3, p. 1354, 2022.',
    '[11] J. Fayyad, M. A. Jaradat, D. Gruyer and H. Najjaran, "Deep Learning Sensor Fusion for Autonomous Vehicle Perception and Localization: A Review," Sensors, vol. 20, no. 15, p. 4220, 2020.',
    '[12] F. Nobis, M. Geisslinger, M. Weber, J. Betz and M. Lienkamp, "A Deep Learning-based Radar and Camera Sensor Fusion Architecture for Object Detection," in Proc. IEEE Sensors Conf., Montreal, QC, Canada, 2019, pp. 1-6.',
    '[13] F. Castanedo, "A Review of Data Fusion Techniques," The Scientific World Journal, vol. 2013, pp. 1-19, 2013.',
    '[14] J. C. Hayward, Near Miss Determination Through Use of a Scale of Danger, Highway Research Record No. 384, Transportation Research Board, Washington, DC, USA, 1972.',
    '[15] J. Jansson and F. Gustafsson, "A Framework and Automotive Application of Collision Avoidance Decision Making," Automatica, vol. 44, no. 9, pp. 2347-2351, 2008.',
    '[16] M. Grieves, Digital Twin: Manufacturing Excellence through Virtual Factory Replication, White Paper, Florida Institute of Technology, Melbourne, FL, USA, 2014.',
    '[17] F. Tao, Q. Qi, A. Liu and A. Kusiak, "Data-Driven Smart Manufacturing," J. Manuf. Syst., vol. 48, pp. 157-169, 2019.',
    '[18] S. Thrun, W. Burgard and D. Fox, Probabilistic Robotics. Cambridge, MA: MIT Press, 2005.',
    '[19] X. Hao, H. Wang, Z. Wang and Y. Lin, "Real-Time Semantic Segmentation for Autonomous Driving on Embedded Hardware," IEEE Trans. Intell. Transp. Syst., vol. 22, no. 7, pp. 4495-4506, 2021.',
    '[20] A. Dosovitskiy, G. Ros, F. Codevilla, A. Lopez and V. Koltun, "CARLA: An Open Urban Driving Simulator," in Proc. 1st Conf. Robot Learning (CoRL), 2017, pp. 1-16.',
    '[21] R. Hassan, "Sensor-Driven Digital Twin Framework for Collision Prevention in Autonomous Systems," GitHub, 2025. [Online]. Available: https://github.com/rh960/Sensor_Driven_Digital_Twin_For_Collison_Prevention_in_Autonomous_Systems',
    '[22] R. Hassan, "Autonomous Systems Project Blog," 2026. [Online]. Available: https://raffayhassan772.wixsite.com/autonomous-systems/blog',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(9.5)

doc.save(r"C:\Sensor Driven Digital Twin For Collision Prevention in Autonomous Systems\Project_Documentation\Chapters_3_4_Final.docx")
print("\nDone! File saved as Chapters_3_4_Final.docx")
print(f"Images successfully embedded: {len(img_cache)}/{len(IMAGES)}")